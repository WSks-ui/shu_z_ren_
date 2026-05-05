# -*- coding: utf-8 -*-
"""
数字人集群 API 端点
提供集群角色查询、讨论流式响应、用户插话等接口
"""

import json
import time
import uuid
from fastapi import APIRouter, Body, HTTPException
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
from io import BytesIO

from py.cluster_roles import get_role_list, get_mode_list, get_all_roles, get_all_role_list, CLUSTER_ROLES, recommend_roles
from py.cluster_orchestrator import (
    create_orchestrator,
    get_orchestrator,
    remove_orchestrator
)
from py.edu_storage import get_cache

router = APIRouter(prefix="/api/cluster", tags=["Digital Human Cluster"])


# ==================== 数据模型 ====================

class ClusterDiscussRequest(BaseModel):
    """集群讨论请求"""
    topic: str = Field(..., description="讨论话题")
    mode: str = Field(default="roundtable", description="讨论模式: roundtable | debate | consultation")
    roles: List[str] = Field(..., description="参与角色ID列表")
    user_input: Optional[str] = Field(default="", description="用户输入（可选）")
    max_rounds: Optional[int] = Field(default=3, description="最大讨论轮次")
    session_id: Optional[str] = Field(default="", description="会话ID")
    continue_from: Optional[str] = Field(default="", description="续聊来源会话ID")


class ClusterInterruptRequest(BaseModel):
    """用户插话请求"""
    session_id: str = Field(..., description="会话ID")
    message: str = Field(..., description="用户消息")


class ClusterTTSRequest(BaseModel):
    """集群 TTS 请求"""
    text: str = Field(..., description="要合成的文本")
    role_id: Optional[str] = Field(default="", description="角色ID（用于选择音色）")


# ==================== API 端点 ====================

@router.get("/roles")
async def get_roles():
    """获取可用角色列表（含自定义角色）"""
    return await get_all_role_list()


@router.get("/modes")
async def get_modes():
    """获取可用讨论模式"""
    return get_mode_list()


@router.post("/discuss/stream")
async def cluster_discuss_stream(request: ClusterDiscussRequest = Body(...)):
    """SSE 流式集群讨论"""

    session_id = request.session_id or f"cluster-{int(time.time())}"

    # 加载历史上下文（续聊模式）
    history = []
    if request.continue_from:
        cache = await get_cache()
        prev_session = await cache.get_cluster_session_detail(request.continue_from)
        if prev_session:
            history = prev_session.get("messages", [])
            # 续聊时继承之前的角色设置
            if not request.roles:
                request.roles = prev_session.get("roles", request.roles)
            # 继承之前的模式
            if not request.mode:
                request.mode = prev_session.get("mode", request.mode)

    # 创建编排器
    orchestrator = await create_orchestrator(
        session_id=session_id,
        mode=request.mode,
        role_ids=request.roles,
        max_rounds=request.max_rounds
    )

    async def generate_stream():
        try:
            async for event in orchestrator.run_discussion(
                topic=request.topic,
                user_input=request.user_input or "",
                history=history,
                session_id=session_id
            ):
                yield event
        except Exception as e:
            print(f"[集群讨论] 错误: {e}")
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"
        finally:
            # 讨论结束后保留编排器一段时间（供插话使用），5分钟后清理
            pass

    return StreamingResponse(
        generate_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@router.post("/interrupt")
async def cluster_interrupt(request: ClusterInterruptRequest = Body(...)):
    """用户插话中断当前讨论"""

    orchestrator = get_orchestrator(request.session_id)
    if not orchestrator:
        raise HTTPException(status_code=404, detail="会话不存在或已结束")

    orchestrator.interrupt(request.message)
    return {"status": "ok", "message": "已插入用户消息"}


@router.get("/affection/{session_id}")
async def get_cluster_affection(session_id: str):
    """获取集群会话的好感度矩阵"""

    orchestrator = get_orchestrator(session_id)
    if not orchestrator:
        raise HTTPException(status_code=404, detail="会话不存在或已结束")

    return {"matrix": orchestrator.get_affection_matrix()}


@router.delete("/session/{session_id}")
async def end_cluster_session(session_id: str):
    """结束集群会话"""

    removed = remove_orchestrator(session_id)
    if not removed:
        raise HTTPException(status_code=404, detail="会话不存在")

    return {"status": "ok", "message": "会话已结束"}


@router.post("/tts")
async def cluster_tts(request: ClusterTTSRequest = Body(...)):
    """集群语音合成 - 使用火山引擎 TTS（回退到 EdgeTTS）"""

    try:
        from py.get_setting import load_settings

        settings = await load_settings()
        tts_settings = settings.get("ttsSettings", {})
        tts_engine = tts_settings.get("engine", "edgetts")

        # 优先使用火山引擎，如果未配置则回退到 EdgeTTS
        if tts_engine == "volcengine" or tts_settings.get("volcAppId"):
            return await _volcengine_tts(request, tts_settings)
        else:
            return await _edgetts_fallback(request, tts_settings)

    except HTTPException:
        raise
    except Exception as e:
        print(f"[集群TTS] 错误: {e}")
        raise HTTPException(status_code=500, detail=str(e))


async def _volcengine_tts(request: ClusterTTSRequest, tts_settings: dict):
    """火山引擎 TTS 合成"""
    import httpx
    import base64

    volc_app_id = tts_settings.get("volcAppId", "")
    volc_access_key = tts_settings.get("volcAccessKey", "")
    volc_resource_id = tts_settings.get("volcResourceId", "volc_tts_release")

    if not volc_app_id or not volc_access_key:
        # 回退到 EdgeTTS
        return await _edgetts_fallback(request, tts_settings)

    # 根据角色选择音色
    voice = _get_role_voice(request.role_id, tts_settings)

    url = "https://openspeech.bytedance.com/api/v3/tts/unidirectional"
    headers = {
        "X-Api-App-Id": volc_app_id,
        "X-Api-Access-Key": volc_access_key,
        "X-Api-Resource-Id": volc_resource_id,
        "Content-Type": "application/json"
    }
    payload = {
        "user": {"uid": "cluster-user"},
        "req_params": {
            "text": request.text,
            "speaker": voice,
            "speed_ratio": 1.0,
            "audio_params": {"format": "mp3", "sample_rate": 24000},
            "additions": json.dumps({"disable_markdown_filter": True})
        }
    }

    collected_audio = bytearray()

    async with httpx.AsyncClient(timeout=60.0) as client:
        async with client.stream("POST", url, headers=headers, json=payload) as response:
            response.raise_for_status()
            async for line in response.aiter_lines():
                if not line:
                    continue
                try:
                    data = json.loads(line)
                    if data.get("code", 0) not in [0, 20000000]:
                        continue
                    if "data" in data and data["data"]:
                        chunk_audio = base64.b64decode(data["data"])
                        collected_audio.extend(chunk_audio)
                except json.JSONDecodeError:
                    continue

    if not collected_audio:
        raise HTTPException(status_code=500, detail="火山引擎TTS合成失败")

    return Response(
        content=bytes(collected_audio),
        media_type="audio/mpeg",
        headers={"Cache-Control": "no-cache"}
    )


async def _edgetts_fallback(request: ClusterTTSRequest, tts_settings: dict):
    """EdgeTTS 回退方案"""
    import edge_tts

    voice = _get_role_voice(request.role_id, tts_settings)
    # 如果 voice 不是 edge tts 格式，使用默认音色
    if not voice.startswith("zh-") and not voice.startswith("en-"):
        voice = "zh-CN-YunxiNeural"

    communicate = edge_tts.Communicate(request.text, voice)
    audio_chunks = []
    async for chunk in communicate.stream():
        if chunk["type"] == "audio":
            audio_chunks.append(chunk["data"])

    if not audio_chunks:
        raise HTTPException(status_code=500, detail="EdgeTTS合成失败")

    return Response(
        content=b"".join(audio_chunks),
        media_type="audio/mpeg",
        headers={"Cache-Control": "no-cache"}
    )


def _get_role_voice(role_id: str, tts_settings: dict = None) -> str:
    """根据角色ID返回对应的音色"""

    # 角色音色映射
    role_voices = {
        "innovator": "BV001_streaming",      # 热情男声
        "skeptic": "BV002_streaming",        # 冷静男声
        "integrator": "BV700_streaming",     # 温和女声
        "practitioner": "BV406_streaming",   # 沉稳男声
    }

    # 如果有用户自定义的火山引擎音色配置，优先使用
    if tts_settings:
        volc_voice = tts_settings.get("volcVoice", "")
        if volc_voice:
            return volc_voice

    return role_voices.get(role_id, "BV001_streaming")


# ==================== 历史记录 API ====================

@router.get("/history")
async def get_cluster_history(mode: str = None, limit: int = 50, offset: int = 0):
    """获取集群讨论历史列表"""

    cache = await get_cache()

    sessions = await cache.get_cluster_sessions(mode=mode, limit=limit, offset=offset)

    # 补充角色名称
    for session in sessions:
        role_names = []
        for role_id in session.get("roles", []):
            role = CLUSTER_ROLES.get(role_id)
            if role:
                role_names.append(role["name"])
            else:
                role_names.append(role_id)
        session["role_names"] = role_names

    return {"sessions": sessions, "total": len(sessions)}


@router.get("/history/{session_id}")
async def get_cluster_history_detail(session_id: str):
    """获取单条集群讨论详情"""

    cache = await get_cache()

    session = await cache.get_cluster_session_detail(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="讨论记录不存在")

    # 补充角色名称和消息中的颜色信息
    role_names = {}
    for role_id in session.get("roles", []):
        role = CLUSTER_ROLES.get(role_id)
        if role:
            role_names[role_id] = {"name": role["name"], "color": role["color"]}
        else:
            role_names[role_id] = {"name": role_id, "color": "#6366f1"}

    session["role_info"] = role_names

    # 为消息添加颜色
    for msg in session.get("messages", []):
        role_id = msg.get("role_id", "")
        if role_id in role_names:
            msg["color"] = role_names[role_id]["color"]

    return session


@router.delete("/history/{session_id}")
async def delete_cluster_history(session_id: str):
    """删除集群讨论记录"""

    cache = await get_cache()

    deleted = await cache.delete_cluster_session(session_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="讨论记录不存在")

    return {"status": "ok", "message": "讨论记录已删除"}


# ==================== 自定义角色 API ====================

class CustomRoleRequest(BaseModel):
    """自定义角色请求"""
    name: str = Field(..., description="角色名称")
    icon: Optional[str] = Field(default="fa-solid fa-user", description="图标")
    color: Optional[str] = Field(default="#6366f1", description="主题色")
    personality: Optional[str] = Field(default="", description="性格描述")
    expertise: Optional[List[str]] = Field(default=[], description="专业领域")
    speaking_style: Optional[str] = Field(default="", description="说话风格")
    system_prompt: str = Field(..., description="系统提示词")
    voice_id: Optional[str] = Field(default="", description="音色ID")


@router.post("/roles/custom")
async def create_custom_role(request: CustomRoleRequest = Body(...)):
    """创建自定义角色"""

    import uuid
    cache = await get_cache()

    role_id = f"custom-{uuid.uuid4().hex[:8]}"
    role_data = {
        "id": role_id,
        "name": request.name,
        "icon": request.icon or "fa-solid fa-user",
        "color": request.color or "#6366f1",
        "personality": request.personality or "",
        "expertise": request.expertise or [],
        "speaking_style": request.speaking_style or "",
        "system_prompt": request.system_prompt,
        "voice_id": request.voice_id or "",
        "created_at": time.time(),
        "updated_at": time.time()
    }

    await cache.add_custom_cluster_role(role_data)
    return {"status": "ok", "role": role_data}


@router.get("/roles/custom/{role_id}")
async def get_custom_role(role_id: str):
    """获取单个自定义角色详情"""

    cache = await get_cache()
    custom_roles = await cache.get_custom_cluster_roles()

    for role in custom_roles:
        if role.get("id") == role_id:
            return {"status": "ok", "role": role}

    raise HTTPException(status_code=404, detail="角色不存在")


@router.put("/roles/custom/{role_id}")
async def update_custom_role(role_id: str, request: CustomRoleRequest = Body(...)):
    """更新自定义角色"""

    cache = await get_cache()

    updates = {
        "name": request.name,
        "icon": request.icon or "fa-solid fa-user",
        "color": request.color or "#6366f1",
        "personality": request.personality or "",
        "expertise": request.expertise or [],
        "speaking_style": request.speaking_style or "",
        "system_prompt": request.system_prompt,
        "voice_id": request.voice_id or ""
    }

    success = await cache.update_custom_cluster_role(role_id, updates)
    if not success:
        raise HTTPException(status_code=404, detail="角色不存在")

    return {"status": "ok", "message": "角色已更新"}


@router.delete("/roles/custom/{role_id}")
async def delete_custom_role(role_id: str):
    """删除自定义角色"""

    cache = await get_cache()

    deleted = await cache.delete_custom_cluster_role(role_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="角色不存在")

    return {"status": "ok", "message": "角色已删除"}


# ==================== 角色记忆 API ====================

@router.get("/roles/{role_id}/memory")
async def get_role_memory(role_id: str):
    """获取角色记忆"""

    cache = await get_cache()

    memories = await cache.get_role_memory(role_id)
    return {"memories": memories}


@router.delete("/roles/{role_id}/memory")
async def clear_role_memory(role_id: str):
    """清除角色记忆"""

    cache = await get_cache()

    cleared = await cache.clear_role_memory(role_id)
    return {"status": "ok", "cleared": cleared}


# ==================== 导出 API ====================

class ClusterExportRequest(BaseModel):
    """导出请求"""
    session_id: str = Field(..., description="会话ID")
    format: str = Field(default="markdown", description="导出格式: markdown | docx")


class ClusterRecommendRequest(BaseModel):
    """角色推荐请求"""
    topic: str = Field(..., description="讨论话题")
    mode: Optional[str] = Field(default="roundtable", description="讨论模式")


@router.post("/roles/recommend")
async def get_role_recommendation(request: ClusterRecommendRequest = Body(...)):
    """根据话题推荐角色组合"""
    return await recommend_roles(request.topic, request.mode or "roundtable")


@router.post("/export")
async def export_cluster_session(request: ClusterExportRequest = Body(...)):
    """导出集群讨论记录"""

    cache = await get_cache()

    session = await cache.get_cluster_session_detail(request.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="讨论记录不存在")

    if request.format == "docx":
        return await _export_cluster_docx(session)
    else:
        return await _export_cluster_markdown(session)


async def _export_cluster_markdown(session: Dict[str, Any]):
    """导出为 Markdown"""
    from io import BytesIO

    lines = []
    lines.append("# 集群讨论记录\n")
    lines.append(f"**话题**: {session['topic']}\n")

    mode_names = {"roundtable": "圆桌讨论", "debate": "正反辩论", "consultation": "导师会诊"}
    lines.append(f"**模式**: {mode_names.get(session['mode'], session['mode'])}\n")

    # 角色名称
    role_names = []
    for role_id in session.get("roles", []):
        role = CLUSTER_ROLES.get(role_id)
        role_names.append(role["name"] if role else role_id)
    lines.append(f"**角色**: {', '.join(role_names)}\n")

    from datetime import datetime
    created = datetime.fromtimestamp(session.get("created_at", 0))
    lines.append(f"**时间**: {created.strftime('%Y-%m-%d %H:%M')}\n")
    lines.append("\n---\n")

    # 按轮次分组
    messages = session.get("messages", [])
    current_round = 0
    for msg in messages:
        round_num = msg.get("round", 1)
        if round_num != current_round:
            current_round = round_num
            lines.append(f"\n## 第{round_num}轮\n")

        role_name = msg.get("role", msg.get("role_id", "未知"))
        content = msg.get("content", "")
        stance = msg.get("stance", "")
        prefix = f"[{stance}方] " if stance else ""
        lines.append(f"\n**{prefix}{role_name}**: {content}\n")

    # 总结
    if session.get("summary"):
        lines.append("\n---\n\n## 讨论总结\n\n")
        lines.append(session["summary"])
        lines.append("\n")

    content = "".join(lines)
    return Response(
        content=content.encode("utf-8"),
        media_type="text/markdown",
        headers={
            "Content-Disposition": f"attachment; filename=cluster-{session['id'][:8]}.md"
        }
    )


async def _export_cluster_docx(session: Dict[str, Any]):
    """导出为 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装，无法导出 Word")

    from io import BytesIO
    from datetime import datetime

    doc = Document()

    # 标题
    title = doc.add_heading("集群讨论记录", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息
    doc.add_paragraph(f"话题：{session['topic']}")
    mode_names = {"roundtable": "圆桌讨论", "debate": "正反辩论", "consultation": "导师会诊"}
    doc.add_paragraph(f"模式：{mode_names.get(session['mode'], session['mode'])}")

    role_names = []
    for role_id in session.get("roles", []):
        role = CLUSTER_ROLES.get(role_id)
        role_names.append(role["name"] if role else role_id)
    doc.add_paragraph(f"角色：{', '.join(role_names)}")

    created = datetime.fromtimestamp(session.get("created_at", 0))
    doc.add_paragraph(f"时间：{created.strftime('%Y-%m-%d %H:%M')}")

    doc.add_paragraph("─" * 40)

    # 对话内容
    messages = session.get("messages", [])
    current_round = 0
    for msg in messages:
        round_num = msg.get("round", 1)
        if round_num != current_round:
            current_round = round_num
            doc.add_heading(f"第{round_num}轮", level=2)

        role_name = msg.get("role", msg.get("role_id", "未知"))
        content = msg.get("content", "")
        stance = msg.get("stance", "")
        prefix = f"[{stance}方] " if stance else ""

        p = doc.add_paragraph()
        run = p.add_run(f"{prefix}{role_name}：")
        run.bold = True
        p.add_run(content)

    # 总结
    if session.get("summary"):
        doc.add_paragraph("─" * 40)
        doc.add_heading("讨论总结", level=2)
        doc.add_paragraph(session["summary"])

    # 导出
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=cluster-{session['id'][:8]}.docx"
        }
    )