# -*- coding: utf-8 -*-
"""
研伴系统 API
包含成长系统、协作记录、情绪状态、对话系统、语音交互等功能的持久化接口
"""

import json
import os
import httpx
import base64
import asyncio
from pathlib import Path
from fastapi import APIRouter, Body, HTTPException, WebSocket, WebSocketDisconnect, Request
from fastapi.responses import StreamingResponse
from typing import Dict, Any, List, Optional, AsyncGenerator
from pydantic import BaseModel, Field
from datetime import datetime

# 模块级别的共享 HTTP 客户端（复用连接池，避免每次请求创建新连接）
_shared_http_client: Optional[httpx.AsyncClient] = None

def get_shared_http_client() -> httpx.AsyncClient:
    """获取共享的 HTTP 客户端，延迟初始化"""
    global _shared_http_client
    if _shared_http_client is None:
        # 配置更健壮的超时和重试设置
        _shared_http_client = httpx.AsyncClient(
            timeout=httpx.Timeout(
                connect=15.0,      # 连接超时 15 秒
                read=60.0,         # 读取超时 60 秒
                write=30.0,        # 写入超时 30 秒
                pool=10.0          # 连接池超时 10 秒
            ),
            limits=httpx.Limits(
                max_connections=100,
                max_keepalive_connections=20,
                keepalive_expiry=30.0
            ),
            # 在 Windows 上启用 HTTP/2 可能导致问题，禁用
            http2=False
        )
    return _shared_http_client

def get_global_http_client():
    """获取全局 HTTP 客户端（复用连接池）"""
    try:
        from server import global_http_client
        if global_http_client is not None:
            return global_http_client
    except Exception:
        pass
    # 回退到模块级别的共享客户端
    return get_shared_http_client()


# 路由前缀
router = APIRouter(prefix="/api/education", tags=["Education Digital Human"])

# 数据存储目录 - 与 edu_storage.py 保持一致
EDUCATION_DATA_DIR = Path(os.path.join(os.path.dirname(__file__), "..", "education_digital_human", "数据"))
EDUCATION_DATA_DIR.mkdir(parents=True, exist_ok=True)

# JSON 文件路径（仅用于健康检查和备份）
GROWTH_DATA_FILE = EDUCATION_DATA_DIR / "growth_data.json"
COLLABORATION_DATA_FILE = EDUCATION_DATA_DIR / "collaboration_data.json"
ACHIEVEMENT_DATA_FILE = EDUCATION_DATA_DIR / "achievement_data.json"
CHAT_HISTORY_FILE = EDUCATION_DATA_DIR / "chat_history.json"

# 使用优化的存储层
from py.edu_storage import get_cache, close_cache

# 教育知识库服务
from py.edu_knowledge_base import (
    get_education_kb,
    search_education_knowledge,
    get_education_context,
    EDUCATION_KB_DIR,
    EDUCATION_VECTOR_DIR,
    preload_education_kb,
    get_education_kb_stats,
    incremental_update_education_kb
)

# 存储层初始化标记
_storage_initialized = False

async def ensure_storage():
    """确保存储层已初始化"""
    global _storage_initialized
    if not _storage_initialized:
        await get_cache()
        _storage_initialized = True

# 技能系统提示词
SKILL_PROMPTS = {
    "research-assistant": """你是"研友"，一个专业的科研助手，专注于帮助研究人员进行学术研究。你的角色是协作者，而非替代者。

## 核心能力
- 文献检索与综述：帮助设计检索策略，选择合适的数据库
- 实验设计支持：帮助明确研究问题和假设，指导选择合适的研究方法
- 数据分析指导：帮助选择合适的统计方法，解释分析结果的含义

## 对话阶段
对话将自然推进以下阶段。请在回复开头用【阶段名】标记当前所处阶段，并根据对话深度主动推进：
- 【选题探索】：帮助用户明确研究方向、缩小范围、找到有价值的研究缺口
- 【假设构建】：引导用户提出可检验的研究假设，明确自变量、因变量和控制变量
- 【实验设计】：协助设计严谨的研究方案，包括样本选择、实验流程、数据收集方法
- 【数据分析】：指导选择统计方法，解读分析结果，讨论研究局限性和改进方向

## 阶段跳转功能
你可以使用特殊指令跳转到指定阶段，格式：[跳转阶段:阶段名]
例如：当讨论已深入到实验设计时，可以输出 [跳转阶段:实验设计] 来更新界面进度。
这在用户主动切换话题或你判断需要回溯到前一阶段时特别有用。

## 追问策略
- 苏格拉底式引导：不直接给答案，而是通过提问帮助用户自己发现答案
- 每次回复至少包含一个开放性问题，引导用户深入思考
- 当用户思路不清晰时，提供 2-3 个选择帮助聚焦
- 在推进下一阶段前，确认用户对当前阶段的理解

## 图片识别能力
你可以识别和处理用户上传的图片，包括：
- 数学公式/手写稿：识别并转换为可编辑的 LaTeX 或文本格式
- 教材插图/实验图：解释原理、步骤或生物学/化学结构
- 笔记截图：整理杂乱笔记，提炼核心知识点
- 数据图表：读取图表数据，分析趋势和结论
- 文献截图：提取关键信息，总结研究内容

## 工作原则
- 引导式协作：通过提问帮助用户澄清思路，提供方法论层面的建议
- 学术诚信：明确标注 AI 生成的内容，提醒用户验证重要信息
- 批判性思维：讨论研究的局限性，提出可能的改进方向

请用专业但友好的语气回答问题，引导用户独立思考。""",

    "literature-review": """你是"研友"，一位文献综述专家，帮助用户进行系统性文献综述和元分析。

## 核心能力
- 检索策略设计：帮助制定 PICO 框架，设计检索词和布尔逻辑
- 文献筛选与管理：协助制定纳入/排除标准，指导使用 PRISMA 流程图
- 质量评估：提供研究质量评估框架，帮助识别潜在偏倚

## 对话阶段
请在回复开头用【阶段名】标记当前阶段，并自然推进：
- 【主题界定】：明确综述的研究问题和范围，确定 PICO 要素
- 【检索策略】：设计系统化的检索方案，选择数据库和关键词
- 【筛选评估】：制定纳入/排除标准，评估文献质量和偏倚风险
- 【综合撰写】：整合研究发现，撰写结构化的综述结论

## 阶段跳转功能
你可以使用特殊指令跳转到指定阶段，格式：[跳转阶段:阶段名]
例如：当用户想跳过检索策略直接讨论筛选标准时，输出 [跳转阶段:筛选评估]

## 追问策略
- 引导用户明确 PICO 各要素后再开始检索
- 每次讨论一篇文献时，追问其与研究问题的关联性
- 提醒用户注意发表偏倚和异质性问题
- 鼓励用户对每篇文献进行批判性评价

## 图片识别能力
你可以识别和处理用户上传的图片，包括：
- 文献截图：提取标题、作者、摘要、关键结论等信息
- 数据图表：读取实验数据图表，辅助数据提取和比较
- 流程图/框架图：理解研究设计和方法流程
- 公式图片：识别数学公式并转换为 LaTeX 格式

## PRISMA 声明提醒
在综述过程中，提醒用户遵循 PRISMA 声明的要求。

请用系统化的方法指导用户完成文献综述。""",

    "paper-writing": """你是"研友"，一位学术论文写作指导专家，帮助用户撰写高质量的学术论文。

## 核心能力
- 论文结构规划：帮助设计论文框架和章节安排
- 写作技巧指导：提供学术写作的语言和风格建议
- 引用格式支持：支持 APA、MLA、Chicago、GB/T 7714、IEEE 等引用格式

## 对话阶段
请在回复开头用【阶段名】标记当前阶段，并自然推进：
- 【论文选题】：明确论文的研究问题和创新点，确定目标期刊/会议
- 【框架搭建】：设计论文整体结构，规划各章节内容和逻辑脉络
- 【内容撰写】：逐章指导写作，包括引言、方法、结果、讨论
- 【润色定稿】：检查学术规范，优化语言表达，完善引用格式

## 阶段跳转功能
你可以使用特殊指令跳转到指定阶段，格式：[跳转阶段:阶段名]
例如：当用户已完成选题，想直接进入写作阶段，输出 [跳转阶段:内容撰写]

## 追问策略
- 先理解论文的研究背景再给写作建议
- 每完成一个章节后，引导用户回顾与整体的逻辑一致性
- 主动指出常见的写作问题（如逻辑跳跃、论据不足）
- 建议用户将长段落拆分为清晰的小节

## 图片识别能力
你可以识别和处理用户上传的图片，包括：
- 数学公式：识别并转换为 LaTeX 格式，可直接用于论文
- 数据图表：帮助描述图表内容，撰写图注说明
- 实验装置图：解释实验原理，辅助撰写方法部分
- 参考文献截图：提取文献信息，生成标准引用格式

## 写作原则
- 学术规范：确保论文符合学术写作规范
- 逻辑清晰：帮助用户理清论证逻辑
- 语言准确：提供专业的学术表达建议

请帮助用户逐步完善论文内容。""",

    "academic-tutoring": """你是"研友"，一位虚拟导师，为学习者提供个性化的学习支持和答疑解惑。

## 核心能力
- 知识讲解：用通俗易懂的方式解释复杂概念
- 学习规划：帮助制定学习计划和目标
- 答疑解惑：回答学习过程中的各种问题

## 对话阶段
请在回复开头用【阶段名】标记当前阶段，并自然推进：
- 【学情诊断】：了解用户的背景、当前水平和学习目标
- 【知识讲解】：用类比和实例解释核心概念，确保用户理解
- 【练习巩固】：设计针对性练习题，帮助用户应用所学知识
- 【总结提升】：归纳知识点，建议下一步学习方向

## 阶段跳转功能
你可以使用特殊指令跳转到指定阶段，格式：[跳转阶段:阶段名]
例如：当用户表示已经理解概念想做练习时，输出 [跳转阶段:练习巩固]

## 追问策略
- 先了解用户的基础水平再调整讲解深度
- 讲解后立即用小问题检验理解程度
- 发现误解时不用否定语气，而是引导用户自己发现矛盾
- 鼓励用户用自己的话复述概念

## 图片识别能力
你可以识别和处理用户上传的图片，包括：
- 数学公式/手写稿：识别并转换为可编辑格式，详细讲解解题步骤
- 教材插图：解释物理、化学、生物等学科的原理和结构
- 笔记截图：整理知识点，帮助复习和记忆
- 错题截图：分析错误原因，提供正确解法和相关知识
- 作业/试卷：批改答案，讲解解题思路

## 教学原则
- 因材施教：根据用户的背景调整讲解深度
- 启发式教学：通过提问引导用户思考
- 耐心细致：确保用户真正理解知识点

请用鼓励和引导的方式帮助用户学习。""",

    "math-assistant": """你是"研友"，一位专业的数学助手，帮助用户解决数学问题、识别数学公式并进行数学推导。

## 核心能力
- 公式识别：识别手写数学公式和印刷公式，转换为 LaTeX 格式
- 分步解题：提供详细的数学解题过程，每一步都标注依据
- 公式推导：进行代数、微积分、线性代数、概率统计等领域的推导
- 错题分析：分析解题过程中的错误，指出错误原因并给出正确解法

## 对话阶段
请在回复开头用【阶段名】标记当前阶段，并自然推进：
- 【问题理解】：确认数学问题的具体内容和已知条件
- 【方法选择】：分析问题类型，选择合适的解题方法
- 【逐步求解】：详细展示每一步推导过程，标注数学依据
- 【验证总结】：验证结果的正确性，总结解题策略和推广

## 阶段跳转功能
你可以使用特殊指令跳转到指定阶段，格式：[跳转阶段:阶段名]
例如：当用户已理解问题想直接看解法时，输出 [跳转阶段:逐步求解]

## 追问策略
- 先确认用户是否理解题目再开始解题
- 关键步骤处暂停，询问用户是否能跟上思路
- 提供多种解法时，先问用户倾向哪种思路
- 解完后引导用户思考该方法的适用范围

## 图片识别能力
你可以识别和处理用户上传的图片，包括：
- 手写数学公式：识别并转换为标准 LaTeX 代码
- 印刷公式：精确识别复杂的数学表达式
- 解题过程截图：逐步检查答案，标注错误步骤
- 几何图形：分析图形关系，辅助几何证明
- 数据表格：提取数据进行统计分析

## 手写公式识别
用户可以通过手写板输入数学公式。当你收到手写公式时：
1. 首先准确识别公式内容
2. 将其转换为标准 LaTeX 格式
3. 解释公式的数学含义
4. 根据用户需求进行后续计算或推导

## 解题原则
- 步骤清晰：每一步都要有明确的数学依据
- 多种方法：优先展示最直观的解法，必要时补充其他方法
- 验证结果：解完后建议用户验证答案的正确性
- 举一反三：总结这类问题的通用解题策略

请用严谨而友好的语气回答数学问题，确保每一步推导都清晰可理解。"""
}


# ==================== 数据模型 ====================

class GrowthStats(BaseModel):
    """成长统计"""
    conversations: int = 0
    papers_read: int = 0
    experiments_designed: int = 0
    papers_written: int = 0
    tutoring_sessions: int = 0
    hours_spent: float = 0.0


class GrowthSystem(BaseModel):
    """成长系统数据"""
    level: int = 1
    exp: int = 0
    totalExp: int = 0
    achievements: List[str] = []
    stats: GrowthStats = GrowthStats()


class CollaborationSession(BaseModel):
    """协作会话记录"""
    id: str
    type: str  # 'paper', 'experiment', 'review', 'tutoring'
    startTime: str
    endTime: Optional[str] = None
    aiContributions: List[Dict[str, Any]] = []
    humanContributions: List[Dict[str, Any]] = []
    summary: Optional[str] = None


class CollaborationRecords(BaseModel):
    """协作记录集合"""
    papers: List[CollaborationSession] = []
    experiments: List[CollaborationSession] = []
    reviews: List[CollaborationSession] = []
    sessions: List[CollaborationSession] = []


class Achievement(BaseModel):
    """成就"""
    id: str
    name: str
    description: str
    icon: str
    unlockedAt: Optional[str] = None


class EmotionState(BaseModel):
    """情绪状态"""
    current: str = "neutral"
    intensity: float = 0.5
    history: List[Dict[str, Any]] = []


class ExportCollaborationRequest(BaseModel):
    """导出协作记录请求"""
    session_id: str
    format: str = "word"
    include_history: bool = True


class ExportAllCollaborationsRequest(BaseModel):
    """导出所有协作记录请求"""
    format: str = "word"
    session_type: Optional[str] = None
    session_ids: Optional[List[str]] = None  # 选择性导出指定会话
    include_chat_history: bool = False  # 是否包含聊天记录详情


class ExportReportRequest(BaseModel):
    """导出学习报告请求"""
    format: str = "word"


# ==================== 辅助函数 ====================

# 最大记录数常量
MAX_COLLABORATION_RECORDS = 500
MAX_CHAT_HISTORY_SESSIONS = 100
MAX_CHAT_MESSAGES_PER_SESSION = 100
SESSION_TIMEOUT_MINUTES = 30  # 会话超时时间（分钟）

# 设置缓存（避免每次请求都访问数据库）
_settings_cache = None
_settings_cache_time = 0
SETTINGS_CACHE_TTL = 5  # 缓存有效期（秒）


async def get_cached_settings() -> Dict[str, Any]:
    """获取缓存的设置，避免频繁访问数据库"""
    global _settings_cache, _settings_cache_time
    import time as time_module

    current_time = time_module.time()

    # 如果缓存有效，直接返回
    if _settings_cache is not None and (current_time - _settings_cache_time) < SETTINGS_CACHE_TTL:
        return _settings_cache

    # 缓存过期或不存在，重新加载
    try:
        from py.get_setting import load_settings
        _settings_cache = await load_settings()
        _settings_cache_time = current_time
        return _settings_cache
    except Exception:
        return {}


def invalidate_settings_cache():
    """使设置缓存失效（保存设置后调用）"""
    global _settings_cache, _settings_cache_time
    _settings_cache = None
    _settings_cache_time = 0


class SessionManager:
    """会话管理器：跟踪活跃会话并自动清理超时会话"""

    def __init__(self, timeout_minutes: int = SESSION_TIMEOUT_MINUTES):
        self.sessions: Dict[str, Dict[str, Any]] = {}
        self.timeout_minutes = timeout_minutes
        self._cleanup_task = None

    def create_session(self, session_id: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """创建新会话"""
        self.sessions[session_id] = {
            "id": session_id,
            "created_at": datetime.now().isoformat(),
            "last_activity": datetime.now().isoformat(),
            "metadata": metadata or {}
        }
        return self.sessions[session_id]

    def get_session(self, session_id: str) -> Optional[Dict[str, Any]]:
        """获取会话并更新活动时间"""
        if session_id in self.sessions:
            self.sessions[session_id]["last_activity"] = datetime.now().isoformat()
            return self.sessions[session_id]
        return None

    def update_activity(self, session_id: str) -> bool:
        """更新会话活动时间"""
        if session_id in self.sessions:
            self.sessions[session_id]["last_activity"] = datetime.now().isoformat()
            return True
        return False

    def remove_session(self, session_id: str) -> bool:
        """移除会话"""
        if session_id in self.sessions:
            del self.sessions[session_id]
            return True
        return False

    def cleanup_inactive_sessions(self) -> int:
        """清理超时会话，返回清理数量"""
        now = datetime.now()
        timeout_threshold = now.timestamp() - (self.timeout_minutes * 60)

        expired_sessions = []
        for session_id, session_data in self.sessions.items():
            last_activity = session_data.get("last_activity", "")
            try:
                last_time = datetime.fromisoformat(last_activity).timestamp()
                if last_time < timeout_threshold:
                    expired_sessions.append(session_id)
            except:
                expired_sessions.append(session_id)

        for session_id in expired_sessions:
            del self.sessions[session_id]

        if expired_sessions:
            print(f"会话管理器: 清理了 {len(expired_sessions)} 个超时会话")

        return len(expired_sessions)

    def get_active_count(self) -> int:
        """获取活跃会话数量"""
        return len(self.sessions)

    def get_all_sessions(self) -> List[Dict[str, Any]]:
        """获取所有会话信息"""
        return list(self.sessions.values())


# 全局会话管理器实例
session_manager = SessionManager()


def _read_json_file_sync(file_path: Path, default: Any = None) -> Any:
    """安全读取 JSON 文件（同步版本）"""
    if not file_path.exists():
        return default
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"读取文件失败 {file_path}: {e}")
        return default


def _write_json_file_sync(file_path: Path, data: Any) -> bool:
    """安全写入 JSON 文件（同步版本）"""
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"写入文件失败 {file_path}: {e}")
        return False


async def _read_json_file(file_path: Path, default: Any = None) -> Any:
    """安全读取 JSON 文件（异步版本，使用线程池避免阻塞）"""
    return await asyncio.to_thread(_read_json_file_sync, file_path, default)


async def _write_json_file(file_path: Path, data: Any) -> bool:
    """安全写入 JSON 文件（异步版本，使用线程池避免阻塞）"""
    return await asyncio.to_thread(_write_json_file_sync, file_path, data)


def _cleanup_collaboration_records(data: dict) -> dict:
    """清理旧协作记录，保留最近 MAX_COLLABORATION_RECORDS 条"""
    total = 0
    max_per_type = MAX_COLLABORATION_RECORDS // 4

    for key in ["papers", "experiments", "reviews", "sessions"]:
        records = data.get(key, [])
        total += len(records)

        if len(records) > max_per_type:
            # 按时间排序，保留最新的
            records.sort(key=lambda x: x.get("startTime", ""), reverse=True)
            data[key] = records[:max_per_type]
            print(f"清理 {key}: 保留 {max_per_type} 条，删除 {len(records) - max_per_type} 条")

    return data


def _cleanup_chat_history(data: dict) -> dict:
    """清理旧对话历史，保留最近 MAX_CHAT_HISTORY_SESSIONS 个会话"""
    sessions = data.get("sessions", {})

    if len(sessions) > MAX_CHAT_HISTORY_SESSIONS:
        # 按最后消息时间排序
        session_list = []
        for sid, msgs in sessions.items():
            last_time = msgs[-1].get("timestamp", "") if msgs else ""
            session_list.append((sid, last_time))

        session_list.sort(key=lambda x: x[1], reverse=True)

        # 保留最新的会话
        keep_ids = {s[0] for s in session_list[:MAX_CHAT_HISTORY_SESSIONS]}
        data["sessions"] = {k: v for k, v in sessions.items() if k in keep_ids}
        print(f"清理对话历史: 保留 {len(data['sessions'])} 个会话")

    return data


# ==================== 成长系统 API ====================

@router.get("/growth")
async def get_growth_data():
    """获取成长系统数据"""
    await ensure_storage()
    cache = await get_cache()
    data = await cache.get("growth", {
        "level": 1,
        "exp": 0,
        "totalExp": 0,
        "achievements": [],
        "stats": {
            "conversations": 0,
            "papers_read": 0,
            "experiments_designed": 0,
            "papers_written": 0,
            "tutoring_sessions": 0,
            "hours_spent": 0.0
        }
    })
    return data


@router.post("/growth")
async def save_growth_data(data: Dict[str, Any] = Body(...)):
    """保存成长系统数据"""
    await ensure_storage()
    cache = await get_cache()
    await cache.set("growth", data)
    return {"status": "success", "message": "成长数据保存成功"}


@router.post("/growth/add_exp")
async def add_experience(amount: int = Body(..., embed=True)):
    """增加经验值（优化的原子操作）"""
    await ensure_storage()
    cache = await get_cache()

    # 获取当前数据
    data = await cache.get("growth", {
        "level": 1,
        "exp": 0,
        "totalExp": 0,
        "achievements": [],
        "stats": {}
    })

    # 使用缓存的原子增量操作
    await cache.increment("growth", "exp", amount)
    await cache.increment("growth", "totalExp", amount)

    # 计算升级
    data = await cache.get("growth")
    exp_for_next = data["level"] * 100
    while data["exp"] >= exp_for_next:
        data["exp"] -= exp_for_next
        data["level"] += 1
        exp_for_next = data["level"] * 100

    await cache.set("growth", data)

    return {
        "status": "success",
        "message": f"获得 {amount} 经验值",
        "newLevel": data["level"],
        "currentExp": data["exp"],
        "expForNext": exp_for_next
    }


@router.post("/growth/update_stats")
async def update_growth_stats(stats: Dict[str, Any] = Body(...)):
    """更新成长统计"""
    await ensure_storage()
    cache = await get_cache()

    # 使用缓存的更新操作
    updates = {}
    for key, value in stats.items():
        if isinstance(value, (int, float)):
            updates[f"stats.{key}"] = value

    for field, value in updates.items():
        await cache.increment("growth", field, int(value) if isinstance(value, (int, float)) else 0)

    data = await cache.get("growth")
    return {"status": "success", "stats": data.get("stats", {})}


# ==================== 成就系统 API ====================

@router.get("/achievements")
async def get_achievements():
    """获取成就列表和解锁状态"""
    await ensure_storage()
    cache = await get_cache()
    data = await cache.get("achievement", {"unlocked": [], "history": []})

    # 定义所有成就
    all_achievements = [
        {"id": "first_chat", "name": "初次对话", "description": "完成第一次对话", "icon": "fa-solid fa-comments"},
        {"id": "paper_reader", "name": "文献读者", "description": "阅读10篇文献", "icon": "fa-solid fa-book"},
        {"id": "experiment_designer", "name": "实验设计师", "description": "设计5个实验方案", "icon": "fa-solid fa-flask"},
        {"id": "paper_writer", "name": "论文写作者", "description": "完成论文写作协助", "icon": "fa-solid fa-pen"},
        {"id": "level_5", "name": "进阶学者", "description": "达到5级", "icon": "fa-solid fa-star"},
        {"id": "level_10", "name": "资深学者", "description": "达到10级", "icon": "fa-solid fa-crown"},
        {"id": "collaboration_master", "name": "协作大师", "description": "完成20次人机协作", "icon": "fa-solid fa-handshake"},
        {"id": "knowledge_seeker", "name": "知识探索者", "description": "使用所有教育技能", "icon": "fa-solid fa-compass"}
    ]

    unlocked_ids = data.get("unlocked", [])

    for ach in all_achievements:
        ach["unlocked"] = ach["id"] in unlocked_ids
        if ach["unlocked"]:
            # 查找解锁时间
            for h in data.get("history", []):
                if h.get("id") == ach["id"]:
                    ach["unlockedAt"] = h.get("unlockedAt")
                    break

    return {"achievements": all_achievements, "unlockedCount": len(unlocked_ids)}


@router.post("/achievements/unlock")
async def unlock_achievement(achievement_id: str = Body(..., embed=True)):
    """解锁成就"""
    await ensure_storage()
    cache = await get_cache()
    data = await cache.get("achievement", {"unlocked": [], "history": []})

    if achievement_id in data.get("unlocked", []):
        return {"status": "already_unlocked", "message": "成就已解锁"}

    data.setdefault("unlocked", []).append(achievement_id)
    data.setdefault("history", []).append({
        "id": achievement_id,
        "unlockedAt": datetime.now().isoformat()
    })

    await cache.set("achievement", data)
    return {"status": "success", "message": f"成就 {achievement_id} 已解锁"}


# ==================== 协作记录 API ====================

@router.get("/collaboration")
async def get_collaboration_data():
    """获取协作记录数据"""
    await ensure_storage()
    cache = await get_cache()
    data = await cache.get("collaboration", {
        "papers": [],
        "experiments": [],
        "reviews": [],
        "sessions": []
    })
    return data


@router.post("/collaboration")
async def save_collaboration_data(data: Dict[str, Any] = Body(...)):
    """保存协作记录数据"""
    await ensure_storage()
    cache = await get_cache()
    await cache.set("collaboration", data)
    return {"status": "success", "message": "协作记录保存成功"}


@router.post("/collaboration/start_session")
async def start_collaboration_session(session_type: str = Body(...), session_id: str = Body(...)):
    """开始新的协作会话"""
    await ensure_storage()
    cache = await get_cache()

    new_session = {
        "id": session_id,
        "type": session_type,
        "startTime": datetime.now().isoformat(),
        "endTime": None,
        "aiContributions": [],
        "humanContributions": [],
        "summary": None
    }

    # 根据类型添加到对应列表
    type_mapping = {
        "paper": "papers",
        "experiment": "experiments",
        "review": "reviews",
        "tutoring": "sessions"
    }

    list_key = type_mapping.get(session_type, "sessions")

    # 获取现有数据并添加新会话
    data = await cache.get("collaboration", {
        "papers": [],
        "experiments": [],
        "reviews": [],
        "sessions": []
    })
    data.setdefault(list_key, []).append(new_session)

    # 同时保存到索引表（支持查询）
    await cache.add_collaboration_session(new_session)

    await cache.set("collaboration", data)
    return {"status": "success", "session": new_session}


@router.post("/collaboration/add_contribution")
async def add_collaboration_contribution(
    session_id: str = Body(...),
    session_type: str = Body(...),
    is_ai: bool = Body(...),
    content: str = Body(...),
    contribution_type: str = Body(default="text")
):
    """向协作会话添加贡献"""
    await ensure_storage()
    cache = await get_cache()

    data = await cache.get("collaboration", {
        "papers": [],
        "experiments": [],
        "reviews": [],
        "sessions": []
    })

    # 找到对应的会话
    type_mapping = {
        "paper": "papers",
        "experiment": "experiments",
        "review": "reviews",
        "tutoring": "sessions"
    }

    list_key = type_mapping.get(session_type, "sessions")

    for session in data.get(list_key, []):
        if session.get("id") == session_id:
            contribution = {
                "content": content,
                "type": contribution_type,
                "timestamp": datetime.now().isoformat()
            }

            if is_ai:
                session.setdefault("aiContributions", []).append(contribution)
            else:
                session.setdefault("humanContributions", []).append(contribution)

            await cache.set("collaboration", data)
            return {"status": "success", "message": "贡献已记录"}

    raise HTTPException(status_code=404, detail="会话不存在")


@router.post("/collaboration/end_session")
async def end_collaboration_session(
    session_id: str = Body(...),
    session_type: str = Body(...),
    summary: str = Body(default="")
):
    """结束协作会话"""
    await ensure_storage()
    cache = await get_cache()

    data = await cache.get("collaboration", {
        "papers": [],
        "experiments": [],
        "reviews": [],
        "sessions": []
    })

    type_mapping = {
        "paper": "papers",
        "experiment": "experiments",
        "review": "reviews",
        "tutoring": "sessions"
    }

    list_key = type_mapping.get(session_type, "sessions")

    for session in data.get(list_key, []):
        if session.get("id") == session_id:
            session["endTime"] = datetime.now().isoformat()
            session["summary"] = summary

            await cache.set("collaboration", data)
            return {"status": "success", "message": "会话已结束"}

    raise HTTPException(status_code=404, detail="会话不存在")


@router.get("/collaboration/stats")
async def get_collaboration_stats():
    """获取协作统计信息"""
    await ensure_storage()
    cache = await get_cache()

    data = await cache.get("collaboration", {
        "papers": [],
        "experiments": [],
        "reviews": [],
        "sessions": []
    })

    stats = {
        "totalSessions": 0,
        "totalAIContributions": 0,
        "totalHumanContributions": 0,
        "byType": {}
    }

    for key in ["papers", "experiments", "reviews", "sessions"]:
        sessions = data.get(key, [])
        stats["totalSessions"] += len(sessions)

        type_stats = {
            "count": len(sessions),
            "aiContributions": 0,
            "humanContributions": 0
        }

        for session in sessions:
            type_stats["aiContributions"] += len(session.get("aiContributions", []))
            type_stats["humanContributions"] += len(session.get("humanContributions", []))
            stats["totalAIContributions"] += len(session.get("aiContributions", []))
            stats["totalHumanContributions"] += len(session.get("humanContributions", []))

        stats["byType"][key] = type_stats

    # 计算协作比例
    total = stats["totalAIContributions"] + stats["totalHumanContributions"]
    if total > 0:
        stats["collaborationRatio"] = {
            "ai": round(stats["totalAIContributions"] / total * 100, 1),
            "human": round(stats["totalHumanContributions"] / total * 100, 1)
        }
    else:
        stats["collaborationRatio"] = {"ai": 0, "human": 0}

    return stats


# ==================== 情绪状态 API ====================

@router.get("/emotion/history")
async def get_emotion_history(limit: int = 100):
    """获取情绪历史记录"""
    await ensure_storage()
    cache = await get_cache()
    growth_data = await cache.get("growth", {})
    history = growth_data.get("emotionHistory", [])[-limit:]
    return {"history": history}


@router.post("/emotion/record")
async def record_emotion(emotion: str = Body(...), intensity: float = Body(...)):
    """记录情绪状态"""
    await ensure_storage()
    cache = await get_cache()

    # 使用缓存的追加操作
    await cache.append_to_list("growth", "emotionHistory", {
        "emotion": emotion,
        "intensity": intensity,
        "timestamp": datetime.now().isoformat()
    }, max_items=500)

    return {"status": "success"}


# ==================== 技能使用统计 API ====================

@router.post("/skills/record_usage")
async def record_skill_usage(skill_id: str = Body(...)):
    """记录技能使用"""
    await ensure_storage()
    cache = await get_cache()

    # 使用缓存的原子增量操作
    await cache.increment("growth", f"skillUsage.{skill_id}")

    data = await cache.get("growth")
    return {"status": "success", "usage": data.get("skillUsage", {})}


@router.get("/skills/usage_stats")
async def get_skill_usage_stats():
    """获取技能使用统计"""
    await ensure_storage()
    cache = await get_cache()
    data = await cache.get("growth", {"skillUsage": {}})
    return {"usageStats": data.get("skillUsage", {})}


# ==================== 书签统计 API ====================

@router.post("/bookmarks/record_save")
async def record_bookmark_save():
    """记录书签保存"""
    await ensure_storage()
    cache = await get_cache()

    # 增加书签统计和经验
    await cache.increment("growth", "stats.bookmarksSaved", 1)
    await cache.increment("growth", "exp", 5)
    await cache.increment("growth", "totalExp", 5)

    data = await cache.get("growth")
    return {"status": "success", "bookmarksSaved": data.get("stats", {}).get("bookmarksSaved", 0)}


@router.post("/bookmarks/record_delete")
async def record_bookmark_delete():
    """记录书签删除（减少统计和经验）"""
    await ensure_storage()
    cache = await get_cache()

    # 减少书签统计（最小为0）
    data = await cache.get("growth", {"stats": {}})
    current = data.get("stats", {}).get("bookmarksSaved", 0)
    if current > 0:
        await cache.increment("growth", "stats.bookmarksSaved", -1)
        # 同时减少经验（最小为0）
        current_exp = data.get("exp", 0)
        if current_exp >= 5:
            await cache.increment("growth", "exp", -5)
        current_total = data.get("totalExp", 0)
        if current_total >= 5:
            await cache.increment("growth", "totalExp", -5)

    return {"status": "success"}


# ==================== 阶段统计 API ====================

@router.post("/stages/record_complete")
async def record_stage_complete(stage_name: str = None):
    """记录阶段完成"""
    await ensure_storage()
    cache = await get_cache()

    # 增加阶段统计和经验
    await cache.increment("growth", "stats.stagesCompleted", 1)
    await cache.increment("growth", "exp", 25)
    await cache.increment("growth", "totalExp", 25)

    data = await cache.get("growth")
    return {
        "status": "success",
        "stagesCompleted": data.get("stats", {}).get("stagesCompleted", 0),
        "stageName": stage_name
    }


# ==================== 健康检查 ====================

@router.get("/health")
async def education_health_check():
    """教育系统健康检查"""
    await ensure_storage()
    cache = await get_cache()

    # 检查数据库连接状态
    db_connected = cache._db_conn is not None

    # 检查知识库状态
    kb_status = "not_initialized"
    kb_doc_count = 0
    try:
        kb = await get_education_kb()
        if kb.vector_store:
            kb_status = "ready"
            kb_doc_count = len(kb.documents)
        elif kb.embeddings:
            kb_status = "configured"
        else:
            kb_status = "no_embeddings"
    except Exception as e:
        kb_status = f"error: {str(e)[:50]}"

    return {
        "status": "ok",
        "dataDir": str(EDUCATION_DATA_DIR),
        "storage": {
            "type": "sqlite",
            "connected": db_connected,
            "cacheKeys": list(cache._cache.keys()) if hasattr(cache, '_cache') else []
        },
        "files": {
            "growth": GROWTH_DATA_FILE.exists(),
            "collaboration": COLLABORATION_DATA_FILE.exists(),
            "achievement": ACHIEVEMENT_DATA_FILE.exists()
        },
        "knowledgeBase": {
            "status": kb_status,
            "documentCount": kb_doc_count,
            "directory": str(EDUCATION_KB_DIR)
        }
    }


# ==================== 对话系统 API ====================

class ChatMessage(BaseModel):
    """对话消息"""
    role: str  # 'user' 或 'assistant'
    content: str
    timestamp: str = Field(default_factory=lambda: datetime.now().isoformat())


class ChatRequest(BaseModel):
    """对话请求"""
    message: str
    skill_id: Optional[str] = None
    session_id: Optional[str] = None
    history: List[Dict[str, str]] = []
    files: Optional[List[str]] = None  # 上传文件的服务器路径列表
    use_fast_model: bool = False  # 强制使用快速模型
    max_tokens: Optional[int] = None  # 覆盖默认max_tokens
    voice_mode: bool = False  # 语音模式：简洁回复
    current_stage: int = 0  # 当前阶段索引
    stage_name: Optional[str] = None  # 当前阶段名称


class ChatResponse(BaseModel):
    """对话响应"""
    response: str
    emotion: Optional[str] = None
    exp_gained: int = 0


@router.post("/chat")
async def education_chat(request: ChatRequest = Body(...)):
    """
    教育对话接口
    自动添加技能系统提示词，调用后端 LLM API
    集成知识库 RAG 检索增强
    """
    import time
    start_time = time.time()

    # 获取系统设置（使用缓存避免频繁访问数据库）
    try:
        t0 = time.time()
        settings = await get_cached_settings()
        print(f"[性能] get_cached_settings 耗时: {time.time() - t0:.2f}s")
    except Exception:
        settings = {}

    # 构建消息
    messages = []

    # 添加技能系统提示词
    if request.skill_id and request.skill_id in SKILL_PROMPTS:
        base_system_prompt = SKILL_PROMPTS[request.skill_id]
    else:
        # 默认教育助手提示词
        base_system_prompt = """你是一位友好的研伴助手，帮助用户学习和研究。你的名字叫"研友"。
请用简洁、专业的语言回答问题，必要时提供学习建议。"""

    # 语音模式：使用简洁提示词，跳过RAG检索
    if request.voice_mode:
        base_system_prompt = """你是语音助手"研友"。你正在用语音和用户对话，必须极其简短地回答。

规则（严格遵守）：
- 回复不超过30个字
- 只说一句话，直接回答
- 不要编号、不要列表、不要Markdown
- 复杂问题回复："这个问题比较复杂，详细解答在聊天窗口里。"
- 禁止说"我可以提供以下支持"之类的话"""
        # 语音模式跳过RAG，加快响应
        rag_context = ""
        rag_sources = []
    else:
        # RAG: 检索知识库获取上下文
        rag_context = ""
        rag_sources = []
        try:
            t_rag = time.time()
            kb = await get_education_kb()
            kb_results = await kb.search(request.message, k=3)

            if kb_results:
                rag_context = kb.get_context_for_chat(kb_results, max_length=1500)
                rag_sources = [r["metadata"].get("file_name", "") for r in kb_results if r.get("metadata", {}).get("file_name")]
                print(f"[RAG] 知识库检索耗时: {time.time() - t_rag:.3f}s, 找到 {len(kb_results)} 条相关内容")
        except Exception as e:
            print(f"[RAG] 知识库检索失败: {e}")

    # 构建最终系统提示词
    if rag_context:
        system_prompt = f"""{base_system_prompt}

## 相关知识库内容

以下是从知识库中检索到的相关内容，请参考这些内容回答用户问题：

{rag_context}

**注意**：
1. 如果知识库内容与用户问题相关，请优先参考知识库内容回答
2. 如果知识库内容不完全匹配，请结合自己的知识补充回答
3. 回答时可以引用知识库来源，格式如「根据《学术写作规范》...」
"""
    else:
        system_prompt = base_system_prompt

    messages.append({
        "role": "system",
        "content": system_prompt
    })

    # 添加历史消息
    for msg in request.history[-20:]:  # 最多保留20条历史
        messages.append({
            "role": msg.get("role", "user"),
            "content": msg.get("content", "")
        })

    # 添加当前消息（支持文件附件）
    user_content = []

    # 处理文件附件
    if request.files:
        import base64
        import mimetypes
        from py.get_setting import UPLOAD_FILES_DIR
        from urllib.parse import urlparse

        for file_path in request.files:
            try:
                local_path = None
                # 将 URL 转换为本地路径
                if file_path.startswith('http://') or file_path.startswith('https://'):
                    parsed = urlparse(file_path)
                    url_path = parsed.path
                    if url_path.startswith('/uploaded_files/'):
                        filename = os.path.basename(url_path)
                        local_path = os.path.join(UPLOAD_FILES_DIR, filename)

                if local_path and os.path.exists(local_path):
                    with open(local_path, 'rb') as f:
                        file_bytes = f.read()
                    mime_type = mimetypes.guess_type(local_path)[0] or 'application/octet-stream'
                elif not file_path.startswith('http://') and not file_path.startswith('https://') and os.path.exists(file_path):
                    with open(file_path, 'rb') as f:
                        file_bytes = f.read()
                    mime_type = mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
                else:
                    print(f"[文件] 文件不存在: {file_path}")
                    continue

                if mime_type.startswith('image/'):
                    b64_data = base64.b64encode(file_bytes).decode('utf-8')
                    user_content.append({
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{b64_data}"}
                    })
                elif mime_type == 'application/pdf':
                    try:
                        from py.load_files import handle_pdf_enhanced
                        pdf_text = await handle_pdf_enhanced(file_bytes)
                        user_content.append({
                            "type": "text",
                            "text": f"\n[PDF文件内容]\n{pdf_text[:3000]}\n[/PDF文件内容]"
                        })
                    except Exception as e:
                        user_content.append({"type": "text", "text": f"\n[PDF文件处理失败: {str(e)}]"})
            except Exception as e:
                print(f"[文件] 处理文件失败 {file_path}: {e}")

    user_content.append({"type": "text", "text": request.message})

    if len(user_content) > 1:
        messages.append({"role": "user", "content": user_content})
    else:
        messages.append({"role": "user", "content": request.message})

    # 获取 API 配置 - 优先使用选中提供商的配置
    api_key = settings.get("api_key", "")
    base_url = settings.get("base_url", "https://api.openai.com/v1")
    model = settings.get("model", "gpt-3.5-turbo")

    selected_provider = settings.get("selectedProvider", None)
    model_providers = settings.get("modelProviders", [])

    # 如果选择了提供商，使用提供商的配置
    if selected_provider and model_providers:
        for provider in model_providers:
            # 统一转换为字符串进行比较
            provider_id = str(provider.get("id", ""))
            selected_id = str(selected_provider)
            if provider_id == selected_id:
                # 使用提供商的配置（数据库字段名）
                if provider.get("apiKey"):
                    api_key = provider.get("apiKey", api_key)
                elif provider.get("api_key"):
                    api_key = provider.get("api_key", api_key)
                if provider.get("url"):
                    base_url = provider.get("url", base_url)
                elif provider.get("base_url"):
                    base_url = provider.get("base_url", base_url)
                if provider.get("modelId"):
                    model = provider.get("modelId", model)
                elif provider.get("model"):
                    model = provider.get("model", model)
                print(f"[教育对话] 使用提供商配置: vendor={provider.get('vendor')}, model={model}, base_url={base_url}")
                break

    if not api_key:
        # 模拟模式下也记录协作
        mock_response = "您好！我是研伴助手。\n\n⚠️ **未配置 API Key**\n\n要启用完整对话功能，请点击页面右上角的「⚙️ 设置」按钮，配置语言模型 API。"
        await update_chat_stats(request.skill_id)
        if request.session_id and request.skill_id:
            await _auto_record_collaboration(
                session_id=request.session_id,
                skill_id=request.skill_id,
                user_message=request.message,
                ai_message=mock_response
            )
        return ChatResponse(
            response=mock_response,
            emotion="neutral",
            exp_gained=5
        )

    def _strip_image_url(msgs):
        """移除消息中的 image_url 内容，仅保留文本"""
        cleaned = []
        for msg in msgs:
            new_msg = dict(msg)
            if isinstance(new_msg.get('content'), list):
                text_items = [item for item in new_msg['content'] if item.get('type') == 'text']
                if text_items:
                    new_msg['content'] = text_items
                else:
                    new_msg['content'] = [{"type": "text", "text": ""}]
            cleaned.append(new_msg)
        return cleaned

    # 调用 LLM API - 使用流式请求以降低延迟
    try:
        t1 = time.time()
        print(f"[性能-语音] 准备调用 LLM API: {base_url}, model={model}, voice_mode={request.voice_mode}")

        request_payload = {
            "model": model,
            "messages": messages,
            "temperature": settings.get("temperature", 0.7),
            "max_tokens": request.max_tokens or settings.get("max_tokens", 2048),
            "stream": True  # 使用流式请求，降低首字节延迟
        }

        # 快速模型：语音模式自动使用快速模型
        if request.use_fast_model:
            t_fast = time.time()
            fast_settings = settings.get("fast", {})
            fast_provider_id = fast_settings.get("selectedProvider")
            print(f"[性能-语音] 快速模型配置: selectedProvider={fast_provider_id}, fast_settings={fast_settings}")
            if fast_provider_id and model_providers:
                for provider in model_providers:
                    if str(provider.get("id", "")) == str(fast_provider_id):
                        if provider.get("apiKey"):
                            api_key = provider.get("apiKey", api_key)
                        elif provider.get("api_key"):
                            api_key = provider.get("api_key", api_key)
                        if provider.get("url"):
                            base_url = provider.get("url", base_url)
                        elif provider.get("base_url"):
                            base_url = provider.get("base_url", base_url)
                        if provider.get("modelId"):
                            model = provider.get("modelId", model)
                        elif provider.get("model"):
                            model = provider.get("model", model)
                        break
            if fast_settings.get("api_key"):
                api_key = fast_settings.get("api_key")
            if fast_settings.get("base_url"):
                base_url = fast_settings.get("base_url")
            if fast_settings.get("model"):
                model = fast_settings.get("model")
            request_payload["model"] = model
            print(f"[性能-语音] 快速模型选择耗时: {(time.time() - t_fast)*1000:.0f}ms, 最终使用模型: {model}")
        else:
            print(f"[性能-语音] 未启用快速模型, 使用默认模型: {model}")

        t_llm_start = time.time()

        # 使用全局客户端（复用连接池）
        global_client = get_global_http_client()
        print(f"[性能-语音] 使用全局客户端: {global_client is not None}")

        assistant_message = ""
        first_chunk_time = None

        async def process_stream(response_stream):
            """处理流式响应，收集内容"""
            nonlocal assistant_message, first_chunk_time
            async for line in response_stream.aiter_lines():
                if not line or line == "data: [DONE]":
                    continue
                if line.startswith("data: "):
                    try:
                        chunk = json.loads(line[6:])
                        delta = chunk.get("choices", [{}])[0].get("delta", {})
                        content = delta.get("content", "")
                        if content:
                            if first_chunk_time is None:
                                first_chunk_time = time.time()
                                print(f"[性能-语音] 首字节延迟: {(first_chunk_time - t_llm_start)*1000:.0f}ms")
                            assistant_message += content
                    except json.JSONDecodeError:
                        continue

        if global_client:
            try:
                async with global_client.stream(
                    "POST",
                    f"{base_url.rstrip('/')}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    },
                    json=request_payload,
                    timeout=30.0  # 语音模式使用较短超时
                ) as response:
                    if response.status_code == 200:
                        await process_stream(response)
                    else:
                        error_body = await response.aread()
                        raise HTTPException(
                            status_code=response.status_code,
                            detail=json.loads(error_body).get("error", {}).get("message", "API 调用失败")
                        )
            except Exception as stream_err:
                print(f"[性能-语音] 全局客户端请求失败: {stream_err}, 回退到临时客户端")
                global_client = None

        if not global_client:
            async with httpx.AsyncClient(timeout=30.0) as client:
                async with client.stream(
                    "POST",
                    f"{base_url.rstrip('/')}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    },
                    json=request_payload
                ) as response:
                    if response.status_code == 200:
                        await process_stream(response)
                    else:
                        error_body = await response.aread()
                        # 如果遇到 image_url 不支持的400错误，移除图片内容后重试
                        if response.status_code == 400:
                            error_text = error_body.decode('utf-8', errors='ignore')
                            if "image_url" in error_text or "image" in error_text.lower():
                                print(f"[教育对话] 模型不支持 image_url，自动移除图片内容后重试")
                                request_payload["messages"] = _strip_image_url(messages)
                                async with client.stream(
                                    "POST",
                                    f"{base_url.rstrip('/')}/chat/completions",
                                    headers={
                                        "Authorization": f"Bearer {api_key}",
                                        "Content-Type": "application/json"
                                    },
                                    json=request_payload
                                ) as retry_response:
                                    await process_stream(retry_response)
                            else:
                                raise HTTPException(
                                    status_code=response.status_code,
                                    detail=json.loads(error_body).get("error", {}).get("message", "API 调用失败")
                                )
                        else:
                            raise HTTPException(
                                status_code=response.status_code,
                                detail=json.loads(error_body).get("error", {}).get("message", "API 调用失败")
                            )

        print(f"[性能-语音] LLM API 调用耗时: {(time.time() - t_llm_start)*1000:.0f}ms")
        print(f"[性能-语音] 回复长度: {len(assistant_message)}字, 内容: {assistant_message[:50]}...")

        # 分析情绪
        t2 = time.time()
        emotion = analyze_emotion(assistant_message)

        # 更新统计
        await update_chat_stats(request.skill_id)
        print(f"[性能] update_chat_stats 耗时: {time.time() - t2:.2f}s")

        # 自动记录协作贡献
        t3 = time.time()
        if request.session_id and request.skill_id:
            await _auto_record_collaboration(
                session_id=request.session_id,
                skill_id=request.skill_id,
                user_message=request.message,
                ai_message=assistant_message
            )
        print(f"[性能] _auto_record_collaboration 耗时: {time.time() - t3:.2f}s")

        print(f"[性能] 总耗时: {time.time() - start_time:.2f}s")
        return ChatResponse(
            response=assistant_message,
            emotion=emotion,
            exp_gained=10
        )

    except httpx.TimeoutException:
        error_response = _get_error_response("llm_timeout")
        raise HTTPException(
            status_code=504,
            detail={
                "code": error_response["code"],
                "message": error_response["message"],
                "detail": error_response["detail"]
            }
        )
    except Exception as e:
        error_type = _classify_error(e)
        error_response = _get_error_response(error_type)
        raise HTTPException(
            status_code=500,
            detail={
                "code": error_response["code"],
                "message": error_response["message"],
                "detail": error_response["detail"],
                "error": str(e)
            }
        )


@router.post("/chat/stream")
async def education_chat_stream(request: ChatRequest = Body(...)):
    """
    教育对话流式接口 (Server-Sent Events)
    实时返回 AI 回复，优化响应体验
    集成知识库 RAG 检索增强
    """
    from fastapi.responses import StreamingResponse

    async def generate_stream() -> AsyncGenerator[str, None]:
        import time
        start_time = time.time()
        full_response = ""

        # 获取系统设置（使用缓存）
        t0 = time.time()
        settings = await get_cached_settings()
        print(f"[性能-流式] get_cached_settings 耗时: {time.time() - t0:.3f}s")

        # 构建消息
        messages = []

        # 语音模式：简洁回复
        voice_mode = getattr(request, 'voice_mode', False)

        # 添加技能系统提示词
        if request.skill_id and request.skill_id in SKILL_PROMPTS:
            base_system_prompt = SKILL_PROMPTS[request.skill_id]
        else:
            base_system_prompt = """你是一位友好的研伴助手，帮助用户学习和研究。你的名字叫"研友"。
请用简洁、专业的语言回答问题，必要时提供学习建议。"""

        # 语音模式使用更简洁的提示词
        if voice_mode:
            base_system_prompt = """你是语音助手"研友"。你正在用语音和用户对话，必须极其简短地回答。

规则（严格遵守）：
- 回复不超过30个字
- 只说一句话，直接回答
- 不要编号、不要列表、不要Markdown
- 复杂问题回复："这个问题比较复杂，详细解答在聊天窗口里。"
- 禁止说"我可以提供以下支持"之类的话"""

        # RAG: 检索知识库获取上下文（仅在向量库已存在时执行）
        rag_context = ""
        rag_sources = []
        try:
            # 检查向量库是否存在，避免每次请求都尝试构建
            index_path = EDUCATION_VECTOR_DIR / "index"

            if index_path.with_suffix(".faiss").exists():
                t_rag = time.time()
                kb = await get_education_kb()
                kb_results = await kb.search(request.message, k=3)

                if kb_results:
                    rag_context = kb.get_context_for_chat(kb_results, max_length=1500)
                    rag_sources = [r["metadata"].get("file_name", "") for r in kb_results if r.get("metadata", {}).get("file_name")]
                    print(f"[RAG-流式] 知识库检索耗时: {time.time() - t_rag:.3f}s, 找到 {len(kb_results)} 条相关内容")
            else:
                print(f"[RAG-流式] 向量库不存在，跳过知识库检索")
        except Exception as e:
            print(f"[RAG-流式] 知识库检索失败: {e}")

        # 构建最终系统提示词
        if rag_context:
            system_prompt = f"""{base_system_prompt}

## 相关知识库内容

以下是从知识库中检索到的相关内容，请参考这些内容回答用户问题：

{rag_context}

**注意**：
1. 如果知识库内容与用户问题相关，请优先参考知识库内容回答
2. 如果知识库内容不完全匹配，请结合自己的知识补充回答
3. 回答时可以引用知识库来源，格式如「根据《学术写作规范》...」
"""
        else:
            system_prompt = base_system_prompt

        # 添加阶段上下文提示（让AI知道当前阶段）
        stage_context = ""
        if request.stage_name:
            stage_context = f"""

## 当前学习阶段
用户当前处于【{request.stage_name}】阶段（第{request.current_stage + 1}阶段）。
请严格按照此阶段的学习目标进行指导：
- 聚焦于当前阶段的核心任务，不要跳过或提前讨论后续阶段的内容
- 如果用户问题超出当前阶段范围，温和地引导回当前阶段
- 只有当用户充分理解当前阶段内容后，才自然推进到下一阶段
- 在回复开头用【{request.stage_name}】标记当前阶段
- 需要跳转阶段时使用指令：[跳转阶段:目标阶段名]"""

        system_prompt += stage_context

        messages.append({
            "role": "system",
            "content": system_prompt
        })

        # 添加历史消息
        for msg in request.history[-10:]:
            messages.append({
                "role": msg.get("role", "user"),
                "content": msg.get("content", "")
            })

        # 添加当前消息（支持文件附件）
        user_content = []

        # 处理文件附件
        if request.files:
            import base64
            import mimetypes
            from py.get_setting import UPLOAD_FILES_DIR
            from urllib.parse import urlparse

            for file_path in request.files:
                try:
                    local_path = None
                    # 将 URL 转换为本地路径
                    if file_path.startswith('http://') or file_path.startswith('https://'):
                        parsed = urlparse(file_path)
                        url_path = parsed.path  # e.g. /uploaded_files/xxx.png
                        if url_path.startswith('/uploaded_files/'):
                            filename = os.path.basename(url_path)
                            local_path = os.path.join(UPLOAD_FILES_DIR, filename)

                    if local_path and os.path.exists(local_path):
                        with open(local_path, 'rb') as f:
                            file_bytes = f.read()
                        mime_type = mimetypes.guess_type(local_path)[0] or 'application/octet-stream'
                    elif not file_path.startswith('http://') and not file_path.startswith('https://') and os.path.exists(file_path):
                        with open(file_path, 'rb') as f:
                            file_bytes = f.read()
                        mime_type = mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
                    else:
                        print(f"[文件] 文件不存在: {file_path}")
                        continue

                    if mime_type.startswith('image/'):
                        # 图片文件：转为 base64 添加到多模态消息
                        b64_data = base64.b64encode(file_bytes).decode('utf-8')
                        user_content.append({
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{b64_data}"
                            }
                        })
                    elif mime_type == 'application/pdf':
                        # PDF 文件：提取文本
                        try:
                            from py.load_files import handle_pdf_enhanced
                            pdf_text = await handle_pdf_enhanced(file_bytes)
                            user_content.append({
                                "type": "text",
                                "text": f"\n[PDF文件内容]\n{pdf_text[:3000]}\n[/PDF文件内容]"
                            })
                        except Exception as e:
                            print(f"[文件] PDF提取失败: {e}")
                            user_content.append({
                                "type": "text",
                                "text": f"\n[PDF文件处理失败: {str(e)}]"
                            })
                    else:
                        # 其他文件类型：尝试读取文本
                        try:
                            text = file_bytes.decode('utf-8', errors='ignore')[:3000]
                            user_content.append({
                                "type": "text",
                                "text": f"\n[文件内容 - {file_path}]\n{text}\n[/文件内容]"
                            })
                        except Exception:
                            pass
                except Exception as e:
                    print(f"[文件] 处理文件失败 {file_path}: {e}")

        # 添加文本消息
        user_content.append({"type": "text", "text": request.message})

        # 如果有文件附件，使用多模态格式；否则使用普通文本格式
        if len(user_content) > 1:
            messages.append({"role": "user", "content": user_content})
        else:
            messages.append({"role": "user", "content": request.message})

        # 获取 API 配置
        api_key = settings.get("api_key", "")
        base_url = settings.get("base_url", "https://api.openai.com/v1")
        model = settings.get("model", "gpt-3.5-turbo")
        temperature = settings.get("temperature", 0.7)
        max_tokens = settings.get("max_tokens", 2048)

        selected_provider = settings.get("selectedProvider", None)
        model_providers = settings.get("modelProviders", [])

        # 如果选择了提供商，使用提供商的配置
        if selected_provider and model_providers:
            for provider in model_providers:
                provider_id = str(provider.get("id", ""))
                selected_id = str(selected_provider)
                if provider_id == selected_id:
                    if provider.get("apiKey"):
                        api_key = provider.get("apiKey", api_key)
                    elif provider.get("api_key"):
                        api_key = provider.get("api_key", api_key)
                    if provider.get("url"):
                        base_url = provider.get("url", base_url)
                    elif provider.get("base_url"):
                        base_url = provider.get("base_url", base_url)
                    if provider.get("modelId"):
                        model = provider.get("modelId", model)
                    elif provider.get("model"):
                        model = provider.get("model", model)
                    break

        # 快速模型逻辑
        fast_settings = settings.get("fast", {})
        use_fast_model = False

        # 如果前端请求强制使用快速模型
        if request.use_fast_model:
            use_fast_model = True
            print("[快速模型] 前端请求强制使用快速模型")
        elif fast_settings.get("enabled", False):
            trigger_mode = fast_settings.get("triggerMode", "conditional")

            if trigger_mode == "always":
                use_fast_model = True
            elif trigger_mode == "conditional":
                # 条件判断
                condition_pass = True
                user_message = request.message or ""

                # 字数限制
                max_len = fast_settings.get("conditionMaxLen", 0)
                if max_len > 0 and len(user_message) > max_len:
                    condition_pass = False

                # 禁止换行
                if condition_pass and fast_settings.get("conditionNoNewline", False):
                    if "\n" in user_message:
                        condition_pass = False

                # 禁止文件/图片 (当前研伴暂不支持文件)
                if condition_pass and fast_settings.get("conditionNoFiles", True):
                    # 如果未来添加文件支持，这里需要检查
                    pass

                if condition_pass:
                    use_fast_model = True

        # 如果触发快速模型，使用快速模型配置
        if use_fast_model:
            fast_provider_id = fast_settings.get("selectedProvider")
            if fast_provider_id and model_providers:
                for provider in model_providers:
                    if str(provider.get("id", "")) == str(fast_provider_id):
                        if provider.get("apiKey"):
                            api_key = provider.get("apiKey", api_key)
                        elif provider.get("api_key"):
                            api_key = provider.get("api_key", api_key)
                        if provider.get("url"):
                            base_url = provider.get("url", base_url)
                        elif provider.get("base_url"):
                            base_url = provider.get("base_url", base_url)
                        if provider.get("modelId"):
                            model = provider.get("modelId", model)
                        elif provider.get("model"):
                            model = provider.get("model", model)
                        break
            # 快速模型也可以有独立的配置覆盖
            if fast_settings.get("api_key"):
                api_key = fast_settings.get("api_key")
            if fast_settings.get("base_url"):
                base_url = fast_settings.get("base_url")
            if fast_settings.get("model"):
                model = fast_settings.get("model")
            if fast_settings.get("temperature") is not None:
                temperature = fast_settings.get("temperature")
            if fast_settings.get("max_tokens") is not None:
                max_tokens = fast_settings.get("max_tokens")

            print(f"[快速模型] 使用快速模型: {model}")

        # 如果前端指定了max_tokens，覆盖默认值
        if request.max_tokens:
            max_tokens = request.max_tokens
            print(f"[快速模型] 使用前端指定的max_tokens: {max_tokens}")

        if not api_key:
            # 模拟模式
            mock_response = "您好！我是研伴助手。\n\n⚠️ **未配置 API Key**\n\n要启用完整对话功能，请配置语言模型 API。"
            yield f"data: {json.dumps({'content': mock_response, 'done': False}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'content': '', 'done': True, 'emotion': 'neutral', 'exp_gained': 5}, ensure_ascii=False)}\n\n"
            return

        # 调用 LLM API (流式) - 优先使用全局客户端（复用连接池）
        global_client = get_global_http_client()
        print(f"[性能-流式] 使用全局客户端: {global_client is not None}")

        # 获取模型选项（是否启用思考）
        model_options = settings.get("modelOptions", {})
        enable_thinking = model_options.get("enableThinking", False)

        # 构建请求体
        request_body = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": True,
            "enable_thinking": enable_thinking  # 根据用户设置决定是否启用思考
        }

        print(f"[性能-流式] enable_thinking: {enable_thinking}")

        try:
            if global_client:
                # 使用全局客户端（连接池已预热）
                async with global_client.stream(
                    "POST",
                    f"{base_url.rstrip('/')}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    },
                    json=request_body,
                    timeout=60.0
                ) as response:
                    async for line in response.aiter_lines():
                        if not line or line == "data: [DONE]":
                            continue
                        if line.startswith("data: "):
                            try:
                                chunk = json.loads(line[6:])
                                delta = chunk.get("choices", [{}])[0].get("delta", {})

                                # 处理思考过程 (reasoning_content) - 实时输出
                                reasoning_content = delta.get("reasoning_content", "")
                                if reasoning_content:
                                    yield f"data: {json.dumps({'reasoning_content': reasoning_content, 'done': False}, ensure_ascii=False)}\n\n"

                                # 处理最终回答 (content)
                                content = delta.get("content", "")
                                if content:
                                    full_response += content
                                    yield f"data: {json.dumps({'content': content, 'done': False}, ensure_ascii=False)}\n\n"
                            except json.JSONDecodeError:
                                continue
            else:
                # 回退：创建临时客户端
                async with httpx.AsyncClient(timeout=60.0) as client:
                    async with client.stream(
                        "POST",
                        f"{base_url.rstrip('/')}/chat/completions",
                        headers={
                            "Authorization": f"Bearer {api_key}",
                            "Content-Type": "application/json"
                        },
                        json=request_body
                    ) as response:
                        async for line in response.aiter_lines():
                            if not line or line == "data: [DONE]":
                                continue
                            if line.startswith("data: "):
                                try:
                                    chunk = json.loads(line[6:])
                                    delta = chunk.get("choices", [{}])[0].get("delta", {})

                                    # 处理思考过程 (reasoning_content) - 实时输出
                                    reasoning_content = delta.get("reasoning_content", "")
                                    if reasoning_content:
                                        yield f"data: {json.dumps({'reasoning_content': reasoning_content, 'done': False}, ensure_ascii=False)}\n\n"

                                    # 处理最终回答 (content)
                                    content = delta.get("content", "")
                                    if content:
                                        full_response += content
                                        yield f"data: {json.dumps({'content': content, 'done': False}, ensure_ascii=False)}\n\n"
                                except json.JSONDecodeError:
                                    continue

            # 流式结束，发送最终元数据
            emotion = analyze_emotion(full_response)
            print(f"[流式对话] 总耗时: {time.time() - start_time:.2f}s, 回复长度: {len(full_response)}")

            # 更新统计和记录协作
            await update_chat_stats(request.skill_id)
            if request.session_id and request.skill_id:
                await _auto_record_collaboration(
                    session_id=request.session_id,
                    skill_id=request.skill_id,
                    user_message=request.message,
                    ai_message=full_response
                )

            # 发送完成信号
            yield f"data: {json.dumps({'content': '', 'done': True, 'emotion': emotion, 'exp_gained': 10}, ensure_ascii=False)}\n\n"

        except httpx.TimeoutException:
            print(f"[流式对话] 超时错误: API 响应超时")
            yield f"data: {json.dumps({'error': 'AI 响应超时，请稍后重试'}, ensure_ascii=False)}\n\n"
        except httpx.ConnectError as e:
            # DNS 解析失败或连接被拒绝
            error_msg = str(e)
            print(f"[流式对话] 连接错误: {error_msg}")
            if "getaddrinfo failed" in error_msg or "11001" in error_msg:
                yield f"data: {json.dumps({'error': '无法连接到 AI 服务，请检查网络连接或 API 地址配置是否正确'}, ensure_ascii=False)}\n\n"
            else:
                yield f"data: {json.dumps({'error': f'连接失败: {error_msg}'}, ensure_ascii=False)}\n\n"
        except httpx.ConnectTimeout:
            print(f"[流式对话] 连接超时: 无法建立与 API 服务器的连接")
            yield f"data: {json.dumps({'error': '连接 AI 服务超时，请检查网络或代理设置'}, ensure_ascii=False)}\n\n"
        except Exception as e:
            error_msg = str(e)
            print(f"[流式对话] 错误: {error_msg}")
            # 检测特定错误类型
            if "getaddrinfo failed" in error_msg or "11001" in error_msg:
                yield f"data: {json.dumps({'error': 'DNS 解析失败，请检查网络连接或 API 地址是否正确'}, ensure_ascii=False)}\n\n"
            else:
                yield f"data: {json.dumps({'error': f'对话失败: {error_msg}'}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        generate_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"  # 禁用 Nginx 缓冲
        }
    )


def analyze_emotion(text: str) -> str:
    """分析文本情绪"""
    emotion_keywords = {
        "happy": ["开心", "高兴", "太好了", "谢谢", "棒", "优秀", "成功", "happy", "great"],
        "thinking": ["为什么", "如何", "怎么", "思考", "分析", "研究", "why", "how"],
        "curious": ["什么是", "能不能", "可以吗", "有趣", "好奇", "what", "interesting"],
        "encouraging": ["加油", "继续", "不错", "很好", "努力", "good", "keep"],
        "neutral": []
    }

    text_lower = text.lower()
    for emotion, keywords in emotion_keywords.items():
        for keyword in keywords:
            if keyword in text_lower:
                return emotion

    return "neutral"


async def update_chat_stats(skill_id: str = None):
    """更新对话统计（使用优化存储层）"""
    await ensure_storage()
    cache = await get_cache()

    # 根据是否使用技能分别统计
    if skill_id:
        # 使用技能对话，统计到技能使用次数
        await cache.increment("growth", "stats.skillUses", 1)
        # 同时更新具体技能的使用次数
        await cache.increment("growth", f"skillUsage.{skill_id}", 1)
    else:
        # 普通对话，统计到对话次数
        await cache.increment("growth", "stats.conversations", 1)

    await cache.increment("growth", "exp", 10)
    await cache.increment("growth", "totalExp", 10)

    # 检查升级（需要读取当前数据）
    data = await cache.get("growth", {
        "level": 1, "exp": 0, "totalExp": 0, "achievements": [], "stats": {},
        "skillUsage": {}
    })

    exp_for_next = data["level"] * 100
    while data["exp"] >= exp_for_next:
        data["exp"] -= exp_for_next
        data["level"] += 1
        exp_for_next = data["level"] * 100

    # 如果有升级，更新数据
    if data["level"] > 1:
        await cache.set("growth", data)


# 技能ID到协作类型的映射
SKILL_TO_COLLAB_TYPE = {
    "research-assistant": "experiment",
    "literature-review": "review",
    "paper-writing": "paper",
    "academic-tutoring": "tutoring",
    "math-assistant": "math"
}

SKILL_TO_TITLE = {
    "research-assistant": "科研助手协作",
    "literature-review": "文献综述协作",
    "paper-writing": "论文写作协作",
    "academic-tutoring": "虚拟导师辅导",
    "math-assistant": "数学解题协作"
}


async def _auto_record_collaboration(
    session_id: str,
    skill_id: str,
    user_message: str,
    ai_message: str
):
    """自动记录协作贡献到协作记录中"""
    await ensure_storage()
    cache = await get_cache()

    collab_type = SKILL_TO_COLLAB_TYPE.get(skill_id, "sessions")
    type_mapping = {
        "paper": "papers",
        "experiment": "experiments",
        "review": "reviews",
        "tutoring": "sessions",
        "math": "sessions"
    }
    list_key = type_mapping.get(collab_type, "sessions")

    data = await cache.get("collaboration", {
        "papers": [], "experiments": [], "reviews": [], "sessions": []
    })

    # 查找已有会话或创建新会话
    found = False
    for session in data.get(list_key, []):
        if session.get("id") == session_id:
            found = True
            # 添加人类贡献
            session.setdefault("humanContributions", []).append({
                "content": user_message[:200],
                "type": "text",
                "timestamp": datetime.now().isoformat()
            })
            # 添加 AI 贡献
            session.setdefault("aiContributions", []).append({
                "content": ai_message[:200],
                "type": "text",
                "timestamp": datetime.now().isoformat()
            })
            # 更新标题（取用户第一条消息的前20个字）
            if not session.get("title") or session["title"] == "":
                session["title"] = user_message[:30] + ("..." if len(user_message) > 30 else "")
            break

    if not found:
        # 创建新会话
        new_session = {
            "id": session_id,
            "type": collab_type,
            "title": user_message[:30] + ("..." if len(user_message) > 30 else ""),
            "startTime": datetime.now().isoformat(),
            "endTime": None,
            "aiContributions": [{
                "content": ai_message[:200],
                "type": "text",
                "timestamp": datetime.now().isoformat()
            }],
            "humanContributions": [{
                "content": user_message[:200],
                "type": "text",
                "timestamp": datetime.now().isoformat()
            }],
            "summary": None
        }
        data.setdefault(list_key, []).append(new_session)
        # 同时保存到索引表（支持查询）
        await cache.add_collaboration_session(new_session)

    # 清理旧记录
    data = _cleanup_collaboration_records(data)
    await cache.set("collaboration", data)


@router.post("/collaboration/generate_summary")
async def generate_collaboration_summary(
    session_id: str = Body(...),
    session_type: str = Body(...)
):
    """为协作会话生成摘要"""
    await ensure_storage()
    cache = await get_cache()

    data = await cache.get("collaboration", {
        "papers": [], "experiments": [], "reviews": [], "sessions": []
    })

    type_mapping = {
        "paper": "papers",
        "experiment": "experiments",
        "review": "reviews",
        "tutoring": "sessions"
    }
    list_key = type_mapping.get(session_type, "sessions")

    for session in data.get(list_key, []):
        if session.get("id") == session_id:
            # 收集所有贡献内容
            all_content = []
            for c in session.get("humanContributions", []):
                all_content.append(f"用户: {c.get('content', '')}")
            for c in session.get("aiContributions", []):
                all_content.append(f"AI: {c.get('content', '')}")

            # 生成简单摘要
            total = len(session.get("humanContributions", [])) + len(session.get("aiContributions", []))
            ai_count = len(session.get("aiContributions", []))
            human_count = len(session.get("humanContributions", []))

            summary = f"共 {total} 次交互（用户 {human_count} 次，AI {ai_count} 次）"
            if all_content:
                # 取第一条用户消息作为摘要开头
                first_msg = session.get("humanContributions", [{}])[0].get("content", "")
                if first_msg:
                    summary = f"关于「{first_msg[:50]}」的协作 - " + summary

            session["summary"] = summary

            # 计算持续时间
            if session.get("startTime"):
                try:
                    start = datetime.fromisoformat(session["startTime"].replace("Z", "+00:00"))
                    end = datetime.now()
                    duration_minutes = int((end - start).total_seconds() / 60)
                    session["duration"] = duration_minutes
                except Exception:
                    pass

            await cache.set("collaboration", data)
            return {"status": "success", "summary": summary}

    raise HTTPException(status_code=404, detail="会话不存在")


@router.get("/chat/history")
async def get_chat_history(session_id: str = None, limit: int = 50):
    """获取对话历史"""
    await ensure_storage()
    cache = await get_cache()
    data = await cache.get("chat_history", {"sessions": {}})

    if session_id:
        return {"history": data.get("sessions", {}).get(session_id, [])}

    # 返回所有会话的最新消息
    all_sessions = data.get("sessions", {})
    result = []
    for sid, messages in all_sessions.items():
        if messages:
            result.append({
                "session_id": sid,
                "last_message": messages[-1] if messages else None,
                "count": len(messages)
            })

    return {"sessions": sorted(result, key=lambda x: x.get("last_message", {}).get("timestamp", ""), reverse=True)[:limit]}


@router.post("/chat/save")
async def save_chat_history(
    session_id: str = Body(...),
    message: Dict[str, Any] = Body(...)
):
    """保存对话消息"""
    await ensure_storage()
    cache = await get_cache()
    data = await cache.get("chat_history", {"sessions": {}})

    data.setdefault("sessions", {})
    data["sessions"].setdefault(session_id, [])

    message["timestamp"] = datetime.now().isoformat()
    data["sessions"][session_id].append(message)

    # 限制每个会话最多保存 MAX_CHAT_MESSAGES_PER_SESSION 条消息
    if len(data["sessions"][session_id]) > MAX_CHAT_MESSAGES_PER_SESSION:
        data["sessions"][session_id] = data["sessions"][session_id][-MAX_CHAT_MESSAGES_PER_SESSION:]

    # 清理旧会话
    data = _cleanup_chat_history(data)

    await cache.set("chat_history", data)
    return {"status": "success"}


@router.delete("/chat/history/{session_id}")
async def delete_chat_history(session_id: str):
    """删除对话历史"""
    await ensure_storage()
    cache = await get_cache()
    data = await cache.get("chat_history", {"sessions": {}})

    if session_id in data.get("sessions", {}):
        del data["sessions"][session_id]
        await cache.set("chat_history", data)
        return {"status": "success"}

    raise HTTPException(status_code=404, detail="会话不存在")


@router.post("/chat/history/update")
async def update_chat_history(
    session_id: str = Body(...),
    messages: List[Dict[str, Any]] = Body(...)
):
    """更新对话历史（用于删除消息等操作）"""
    await ensure_storage()
    cache = await get_cache()
    data = await cache.get("chat_history", {"sessions": {}})

    data.setdefault("sessions", {})
    data["sessions"][session_id] = messages

    # 清理旧会话
    data = _cleanup_chat_history(data)

    await cache.set("chat_history", data)
    return {"status": "success"}


# ==================== 会话管理 API ====================

@router.post("/session/create")
async def create_session(
    session_id: str = Body(...),
    skill_id: str = Body(None),
    digital_human_type: str = Body("vrm")
):
    """创建新会话"""
    metadata = {
        "skill_id": skill_id,
        "digital_human_type": digital_human_type
    }
    session = session_manager.create_session(session_id, metadata)
    return {"status": "success", "session": session}


@router.get("/session/{session_id}")
async def get_session(session_id: str):
    """获取会话信息"""
    session = session_manager.get_session(session_id)
    if session:
        return {"status": "success", "session": session}
    return {"status": "not_found", "session": None}


@router.post("/session/{session_id}/activity")
async def update_session_activity(session_id: str):
    """更新会话活动时间"""
    if session_manager.update_activity(session_id):
        return {"status": "success"}
    return {"status": "not_found"}


@router.delete("/session/{session_id}")
async def delete_session(session_id: str):
    """删除会话"""
    if session_manager.remove_session(session_id):
        return {"status": "success"}
    return {"status": "not_found"}


@router.post("/session/cleanup")
async def cleanup_sessions():
    """手动触发清理超时会话"""
    cleaned = session_manager.cleanup_inactive_sessions()
    return {"status": "success", "cleaned_count": cleaned}


@router.get("/sessions/active")
async def get_active_sessions():
    """获取所有活跃会话"""
    sessions = session_manager.get_all_sessions()
    return {
        "status": "success",
        "count": len(sessions),
        "sessions": sessions
    }


# ==================== 技能提示词 API ====================

@router.get("/skills/prompts")
async def get_skill_prompts():
    """获取所有技能的系统提示词"""
    return {"prompts": SKILL_PROMPTS}


@router.get("/skills/{skill_id}/prompt")
async def get_skill_prompt(skill_id: str):
    """获取指定技能的系统提示词"""
    if skill_id in SKILL_PROMPTS:
        return {"skill_id": skill_id, "prompt": SKILL_PROMPTS[skill_id]}
    raise HTTPException(status_code=404, detail="技能不存在")


# ==================== API 设置 ====================

# 允许访问设置 API 的来源列表
ALLOWED_ORIGINS = [
    "http://127.0.0.1",
    "http://localhost",
    "http://0.0.0.0",
]


def _validate_request_origin(request) -> bool:
    """
    验证请求来源是否合法

    Args:
        request: FastAPI 请求对象

    Returns:
        是否为合法来源
    """
    from fastapi import Request

    # 获取来源信息
    origin = request.headers.get("origin", "")
    referer = request.headers.get("referer", "")

    # 检查是否为本地访问
    client_host = getattr(request, "client", None)
    if client_host:
        client_ip = getattr(client_host, "host", "")
        if client_ip in ["127.0.0.1", "::1", "localhost"]:
            return True

    # 检查 Origin 头
    if origin:
        for allowed in ALLOWED_ORIGINS:
            if origin.startswith(allowed):
                return True

    # 检查 Referer 头
    if referer:
        for allowed in ALLOWED_ORIGINS:
            if referer.startswith(allowed):
                return True

    # 如果没有 Origin 和 Referer，可能是直接 API 调用（允许本地开发）
    if not origin and not referer:
        return True

    return False


def _mask_api_key(api_key: str) -> str:
    """
    隐藏 API Key，只显示后4位

    Args:
        api_key: 原始 API Key

    Returns:
        隐藏后的 API Key
    """
    if not api_key or len(api_key) <= 4:
        return "****" if api_key else ""
    return "****" + api_key[-4:]


@router.get("/settings")
async def get_api_settings(request: Request):
    """获取 API 设置"""
    try:
        from py.get_setting import load_settings
        settings = await load_settings()

        # 获取根级别配置
        api_key = settings.get("api_key", "")
        base_url = settings.get("base_url", "")
        model = settings.get("model", "")
        selected_provider = settings.get("selectedProvider", None)
        model_providers = settings.get("modelProviders", [])

        # 调试日志
        print(f"[研伴设置] selectedProvider: {selected_provider} (type: {type(selected_provider)})")
        print(f"[研伴设置] modelProviders 数量: {len(model_providers)}")
        for p in model_providers:
            print(f"  - Provider ID: {p.get('id')} (type: {type(p.get('id'))}), vendor: {p.get('vendor')}, has apiKey: {bool(p.get('apiKey'))}")

        # 如果选择了提供商，使用提供商的配置
        provider_info = None
        if selected_provider and model_providers:
            for provider in model_providers:
                # 统一转换为字符串进行比较，避免类型不匹配问题
                provider_id = str(provider.get("id", ""))
                selected_id = str(selected_provider)
                if provider_id == selected_id:
                    # 数据库字段映射：vendor->name, url->base_url, apiKey->api_key, modelId->model
                    vendor_name = provider.get("vendor", selected_provider)
                    # 映射 vendor 到友好名称
                    vendor_display_names = {
                        "aliyun": "阿里云百炼",
                        "openai": "OpenAI",
                        "anthropic": "Anthropic",
                        "deepseek": "DeepSeek",
                        "zhipu": "智谱AI",
                        "moonshot": "Moonshot",
                        "baidu": "百度文心",
                        "tencent": "腾讯混元",
                    }
                    display_name = vendor_display_names.get(vendor_name, vendor_name)

                    provider_info = {
                        "name": display_name,
                        "type": provider.get("type", "custom"),
                        "model": provider.get("modelId", provider.get("model", "")),
                        "base_url": provider.get("url", provider.get("base_url", "")),
                    }
                    # 如果提供商有配置，优先使用提供商的值
                    if provider.get("apiKey"):
                        api_key = provider.get("apiKey", api_key)
                    elif provider.get("api_key"):
                        api_key = provider.get("api_key", api_key)
                    if provider.get("url"):
                        base_url = provider.get("url", base_url)
                    elif provider.get("base_url"):
                        base_url = provider.get("base_url", base_url)
                    if provider.get("modelId"):
                        model = provider.get("modelId", model)
                    elif provider.get("model"):
                        model = provider.get("model", model)
                    break

        # 构建显示名称
        display_model = model
        if provider_info:
            display_model = f"[{provider_info['name']}] {model}"

        # 转换 modelProviders 列表，统一字段名
        vendor_display_names = {
            "aliyun": "阿里云百炼",
            "openai": "OpenAI",
            "anthropic": "Anthropic",
            "deepseek": "DeepSeek",
            "zhipu": "智谱AI",
            "moonshot": "Moonshot",
            "baidu": "百度文心",
            "tencent": "腾讯混元",
        }

        transformed_providers = []
        for p in model_providers:
            vendor_name = p.get("vendor", p.get("name", str(p.get("id", ""))))
            display_name = vendor_display_names.get(vendor_name, p.get("name", vendor_name))

            transformed_providers.append({
                "id": p.get("id"),
                "name": display_name,
                "type": p.get("type", "custom"),
                "model": p.get("modelId", p.get("model", "")),
                "base_url": p.get("url", p.get("base_url", "")),
                "vendor": vendor_name,
            })

        # 获取快速模型设置
        fast_settings = settings.get("fast", {})

        # 获取模型选项设置
        model_options = settings.get("modelOptions", {"enableThinking": False})

        return {
            "api_key": _mask_api_key(api_key),  # 隐藏 API Key
            "api_key_configured": bool(api_key),  # 是否已配置
            "base_url": base_url,
            "model": display_model,
            "raw_model": model,  # 原始模型名
            "selectedProvider": selected_provider,
            "providerInfo": provider_info,  # 提供商信息
            "temperature": settings.get("temperature", 0.7),
            "max_tokens": settings.get("max_tokens", 2048),
            "modelProviders": transformed_providers,  # 转换后的模型提供商列表
            "fast": fast_settings,  # 快速模型设置
            "modelOptions": model_options,  # 模型选项
            "formulaOcr": settings.get("formulaOcrSettings", {"enabled": True, "api_key": "", "model": "standard"})
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"加载设置失败: {str(e)}")


@router.post("/settings")
async def save_api_settings(settings: Dict[str, Any] = Body(...), request: Request = None):
    """保存 API 设置（包括快速模型设置）"""
    try:
        # 验证请求来源
        if request and not _validate_request_origin(request):
            raise HTTPException(
                status_code=403,
                detail="禁止访问：设置 API 仅允许从本地访问"
            )

        from py.get_setting import load_settings, save_settings

        # 加载现有设置
        current_settings = await load_settings()

        # 更新基础设置
        if "api_key" in settings:
            current_settings["api_key"] = settings["api_key"]
        if "base_url" in settings:
            current_settings["base_url"] = settings["base_url"]
        if "model" in settings:
            current_settings["model"] = settings["model"]
        if "temperature" in settings:
            current_settings["temperature"] = settings["temperature"]
        if "max_tokens" in settings:
            current_settings["max_tokens"] = settings["max_tokens"]

        # 更新快速模型设置
        if "fast" in settings:
            current_settings["fast"] = settings["fast"]
            print(f"[研伴设置] 保存快速模型配置: {settings['fast']}")

        # 更新模型选项设置
        if "modelOptions" in settings:
            current_settings["modelOptions"] = settings["modelOptions"]
            print(f"[研伴设置] 保存模型选项: {settings['modelOptions']}")

        # 更新公式识别设置
        if "formulaOcr" in settings:
            current_settings["formulaOcrSettings"] = settings["formulaOcr"]
            print(f"[研伴设置] 保存公式识别设置")

        # 保存设置
        await save_settings(current_settings)

        # 使缓存失效
        invalidate_settings_cache()

        return {"status": "success", "message": "设置已保存"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"保存设置失败: {str(e)}")


# ==================== 综合统计 API ====================

@router.get("/dashboard")
async def get_dashboard():
    """获取仪表盘数据"""
    await ensure_storage()
    cache = await get_cache()

    growth_data = await cache.get("growth", {
        "level": 1, "exp": 0, "totalExp": 0, "achievements": [], "stats": {}
    })

    collab_data = await cache.get("collaboration", {
        "papers": [], "experiments": [], "reviews": [], "sessions": []
    })

    # 计算协作统计
    total_sessions = sum(len(collab_data.get(k, [])) for k in ["papers", "experiments", "reviews", "sessions"])

    return {
        "level": growth_data.get("level", 1),
        "exp": growth_data.get("exp", 0),
        "totalExp": growth_data.get("totalExp", 0),
        "achievements": len(growth_data.get("achievements", [])),
        "stats": growth_data.get("stats", {}),
        "collaborationSessions": total_sessions,
        "skillUsage": growth_data.get("skillUsage", {})
    }


# ==================== 语音交互 API ====================

class VoiceChatRequest(BaseModel):
    """语音对话请求"""
    audio_data: str  # base64 编码的音频数据
    skill_id: Optional[str] = None
    session_id: Optional[str] = None
    language: str = "zh"
    digital_human_type: str = "vrm"  # "vrm" 或 "tencent" 或 "xingyun"
    tencent_session_id: Optional[str] = None  # 腾讯数智人会话 ID
    history: List[Dict[str, str]] = []  # 历史对话消息
    frontend_drive: bool = False  # 前端自行驱动数字人播报，后端跳过TTS合成


class VoiceChatResponse(BaseModel):
    """语音对话响应"""
    text: str  # 识别的文本
    response: str  # AI回复文本
    audio_data: Optional[str] = None  # base64 编码的音频响应（仅 VRM 模式）
    emotion: Optional[str] = None
    exp_gained: int = 0
    digital_human_driven: bool = False  # 是否已驱动数字人播报


@router.post("/voice/chat")
async def voice_chat(request: VoiceChatRequest = Body(...)):
    """
    语音对话接口
    1. 接收音频数据，使用 Sherpa ASR 进行语音识别
    2. 调用 LLM 生成回复
    3. 根据数字人类型选择输出方式：
       - 腾讯数智人：直接驱动数字人播报，不返回 audio_data
       - VRM 数字人：使用 TTS 合成音频返回
    """
    try:
        # 1. 解码音频数据
        audio_bytes = base64.b64decode(request.audio_data)

        # 2. 语音识别 (Sherpa ASR)
        recognized_text = await _recognize_speech(audio_bytes, request.language)

        if not recognized_text or not recognized_text.strip():
            return VoiceChatResponse(
                text="",
                response="抱歉，我没有听清，请再说一次。",
                exp_gained=0
            )

        # 3. 调用 LLM 获取回复
        chat_response = await _get_llm_response(
            message=recognized_text,
            skill_id=request.skill_id,
            session_id=request.session_id,
            history=request.history
        )

        response_text = chat_response["response"]
        audio_data = None
        digital_human_driven = False

        # 4. 根据数字人类型处理语音输出
        if request.frontend_drive:
            # 前端自行驱动数字人（如星云），后端跳过TTS合成
            digital_human_driven = True
        elif request.digital_human_type == "tencent" and request.tencent_session_id:
            # 腾讯数智人：直接驱动数字人播报，不生成独立音频
            try:
                await _drive_tencent_digital_human(
                    session_id=request.tencent_session_id,
                    text=response_text
                )
                digital_human_driven = True
            except Exception as e:
                print(f"驱动腾讯数智人失败: {e}")
                # 失败时降级到 TTS
                audio_data = await _synthesize_speech(response_text, request.language)
        else:
            # VRM 或无数字人：使用 TTS 合成音频
            audio_data = await _synthesize_speech(response_text, request.language)

        # 5. 保存对话历史
        if request.session_id:
            await _save_voice_chat_history(
                session_id=request.session_id,
                user_message=recognized_text,
                assistant_message=response_text,
                skill_id=request.skill_id
            )

        return VoiceChatResponse(
            text=recognized_text,
            response=response_text,
            audio_data=audio_data,
            emotion=chat_response.get("emotion"),
            exp_gained=chat_response.get("exp_gained", 10),
            digital_human_driven=digital_human_driven
        )

    except HTTPException:
        raise
    except Exception as e:
        error_str = str(e)
        # 区分 ASR 模型不可用与其他错误
        if "ASR 模型未就绪" in error_str or "sherpa" in error_str.lower():
            raise HTTPException(
                status_code=503,
                detail={
                    "code": "ASR_UNAVAILABLE",
                    "message": "语音识别服务暂不可用",
                    "detail": "Sherpa ASR 模型未加载，请检查模型文件是否已下载。您可以使用文字输入进行对话。",
                    "error": error_str
                }
            )
        error_type = _classify_error(e)
        error_response = _get_error_response(error_type)
        raise HTTPException(
            status_code=500,
            detail={
                "code": error_response["code"],
                "message": error_response["message"],
                "detail": error_response["detail"],
                "error": error_str
            }
        )


@router.post("/voice/recognize")
async def voice_recognize(request: Dict[str, Any] = Body(...)):
    """
    单独的语音识别接口
    """
    try:
        audio_data = request.get("audio_data", "")
        language = request.get("language", "zh")
        audio_bytes = base64.b64decode(audio_data)
        text = await _recognize_speech(audio_bytes, language)
        return {"text": text, "success": True}
    except Exception as e:
        return {"text": "", "success": False, "error": str(e)}


@router.post("/voice/synthesize")
async def voice_synthesize(request: Dict[str, Any] = Body(...)):
    """
    单独的语音合成接口
    """
    try:
        text = request.get("text", "")
        language = request.get("language", "zh")
        engine = request.get("engine", "edge")
        voice_id = request.get("voice_id", "default")
        audio_bytes = await _synthesize_speech(text, language, engine, voice_id)
        return {"audio_data": audio_bytes, "success": True}
    except Exception as e:
        return {"audio_data": None, "success": False, "error": str(e)}


@router.websocket("/ws/voice")
async def websocket_voice_chat(websocket: WebSocket):
    """
    WebSocket 语音对话接口
    实现实时语音交互

    消息协议:
    - 客户端 -> 服务端:
      { "type": "audio", "data": "base64_audio", "language": "zh" }
      { "type": "text", "data": "文本消息", "skill_id": "...", "session_id": "..." }
      { "type": "config", "skill_id": "...", "session_id": "...", "digital_human_type": "vrm", "tencent_session_id": "..." }

    - 服务端 -> 客户端:
      { "type": "recognized", "text": "识别的文本" }
      { "type": "response", "text": "AI回复" }
      { "type": "audio", "data": "base64_audio" }  // 仅 VRM 模式
      { "type": "digital_human_driven", "success": true }  // 腾讯模式
      { "type": "emotion", "emotion": "happy" }
      { "type": "error", "message": "错误信息" }
    """
    await websocket.accept()

    session_config = {
        "skill_id": None,
        "session_id": "voice_" + datetime.now().strftime("%Y%m%d_%H%M%S"),
        "language": "zh",
        "digital_human_type": "vrm",
        "tencent_session_id": None,
        "history": []  # 对话历史，最多保留20条
    }

    try:
        while True:
            # 接收消息
            data = await websocket.receive_text()
            message = json.loads(data)
            msg_type = message.get("type")

            if msg_type == "config":
                # 更新会话配置
                session_config["skill_id"] = message.get("skill_id")
                session_config["session_id"] = message.get("session_id", session_config["session_id"])
                session_config["language"] = message.get("language", "zh")
                session_config["digital_human_type"] = message.get("digital_human_type", "vrm")
                session_config["tencent_session_id"] = message.get("tencent_session_id")

                # 创建或更新会话管理器中的会话
                session_manager.create_session(
                    session_config["session_id"],
                    {
                        "skill_id": session_config["skill_id"],
                        "digital_human_type": session_config["digital_human_type"],
                        "source": "websocket"
                    }
                )

                await websocket.send_json({"type": "config", "status": "ok"})

            elif msg_type == "audio":
                # 更新会话活动时间
                session_manager.update_activity(session_config["session_id"])

                # 语音识别
                audio_data = message.get("data", "")
                language = message.get("language", session_config["language"])

                try:
                    audio_bytes = base64.b64decode(audio_data)
                    recognized_text = await _recognize_speech(audio_bytes, language)

                    await websocket.send_json({
                        "type": "recognized",
                        "text": recognized_text
                    })

                    if recognized_text and recognized_text.strip():
                        # 获取 AI 回复（传递历史消息）
                        response = await _get_llm_response(
                            message=recognized_text,
                            skill_id=session_config["skill_id"],
                            session_id=session_config["session_id"],
                            history=session_config["history"]
                        )

                        # 更新对话历史
                        session_config["history"].append({"role": "user", "content": recognized_text})
                        session_config["history"].append({"role": "assistant", "content": response["response"]})
                        session_config["history"] = session_config["history"][-20:]

                        # 发送回复文本
                        await websocket.send_json({
                            "type": "response",
                            "text": response["response"]
                        })

                        # 发送情绪
                        if response.get("emotion"):
                            await websocket.send_json({
                                "type": "emotion",
                                "emotion": response["emotion"]
                            })

                        # 根据数字人类型处理语音输出
                        if session_config["digital_human_type"] == "tencent" and session_config["tencent_session_id"]:
                            # 腾讯数智人：直接驱动数字人播报
                            success = await _drive_tencent_digital_human(
                                session_id=session_config["tencent_session_id"],
                                text=response["response"]
                            )
                            await websocket.send_json({
                                "type": "digital_human_driven",
                                "success": success
                            })
                        else:
                            # VRM 或无数字人：合成音频返回
                            audio_response = await _synthesize_speech(response["response"], language)
                            await websocket.send_json({
                                "type": "audio",
                                "data": audio_response
                            })

                        # 保存历史
                        await _save_voice_chat_history(
                            session_id=session_config["session_id"],
                            user_message=recognized_text,
                            assistant_message=response["response"],
                            skill_id=session_config["skill_id"]
                        )

                except Exception as e:
                    error_type = _classify_error(e)
                    error_response = _get_error_response(error_type, str(e))
                    await websocket.send_json({
                        "type": "error",
                        "code": error_response["code"],
                        "message": error_response["message"],
                        "detail": error_response["detail"]
                    })

            elif msg_type == "text":
                # 更新会话活动时间
                session_manager.update_activity(session_config["session_id"])

                # 直接处理文本消息
                text_data = message.get("data", "")

                try:
                    response = await _get_llm_response(
                        message=text_data,
                        skill_id=session_config["skill_id"],
                        session_id=session_config["session_id"],
                        history=session_config["history"]
                    )

                    # 更新对话历史
                    session_config["history"].append({"role": "user", "content": text_data})
                    session_config["history"].append({"role": "assistant", "content": response["response"]})
                    session_config["history"] = session_config["history"][-20:]

                    await websocket.send_json({
                        "type": "response",
                        "text": response["response"]
                    })

                    # 根据数字人类型处理语音输出
                    if session_config["digital_human_type"] == "tencent" and session_config["tencent_session_id"]:
                        # 腾讯数智人：直接驱动数字人播报
                        success = await _drive_tencent_digital_human(
                            session_id=session_config["tencent_session_id"],
                            text=response["response"]
                        )
                        await websocket.send_json({
                            "type": "digital_human_driven",
                            "success": success
                        })
                    else:
                        # VRM 或无数字人：合成音频返回
                        audio_response = await _synthesize_speech(
                            response["response"],
                            session_config["language"]
                        )
                        await websocket.send_json({
                            "type": "audio",
                            "data": audio_response
                        })

                except Exception as e:
                    error_type = _classify_error(e)
                    error_response = _get_error_response(error_type, str(e))
                    await websocket.send_json({
                        "type": "error",
                        "code": error_response["code"],
                        "message": error_response["message"],
                        "detail": error_response["detail"]
                    })

    except WebSocketDisconnect:
        session_manager.remove_session(session_config["session_id"])
        print(f"WebSocket 语音连接断开: {session_config['session_id']}")
    except Exception as e:
        session_manager.remove_session(session_config["session_id"])
        print(f"WebSocket 错误: {e}")


# ==================== 辅助函数 ====================

# 错误消息映射表
ERROR_MESSAGES = {
    # 语音相关
    "speech_recognition_failed": {
        "code": "SPEECH_001",
        "message": "语音识别失败，请确保麦克风正常工作",
        "detail": "可能原因：麦克风未授权、环境噪音过大、或语音过短"
    },
    "speech_synthesis_failed": {
        "code": "SPEECH_002",
        "message": "语音合成失败，请稍后重试",
        "detail": "TTS 服务暂时不可用"
    },

    # API 相关
    "llm_timeout": {
        "code": "API_001",
        "message": "AI 响应超时，请稍后重试",
        "detail": "服务繁忙，请稍后再试"
    },
    "llm_error": {
        "code": "API_002",
        "message": "AI 服务暂时不可用",
        "detail": "请检查网络连接或稍后重试"
    },
    "dns_resolution_failed": {
        "code": "API_004",
        "message": "无法连接到 AI 服务",
        "detail": "DNS 解析失败，请检查网络连接或 API 地址配置是否正确"
    },
    "api_rate_limit": {
        "code": "API_003",
        "message": "请求过于频繁，请稍后再试",
        "detail": "已达到 API 调用限制"
    },

    # 会话相关
    "session_not_found": {
        "code": "SESSION_001",
        "message": "会话已过期，请重新开始",
        "detail": "会话可能已被清理或不存在"
    },
    "session_expired": {
        "code": "SESSION_002",
        "message": "会话已超时，请刷新页面",
        "detail": "长时间未活动，会话已自动关闭"
    },

    # 数据相关
    "data_save_failed": {
        "code": "DATA_001",
        "message": "数据保存失败，请重试",
        "detail": "存储服务暂时不可用"
    },
    "data_load_failed": {
        "code": "DATA_002",
        "message": "数据加载失败，请刷新页面",
        "detail": "无法读取历史数据"
    },

    # 数字人相关
    "digital_human_connect_failed": {
        "code": "DH_001",
        "message": "数字人连接失败，请检查网络",
        "detail": "无法连接到数字人服务"
    },

    # 通用错误
    "unknown_error": {
        "code": "UNKNOWN_001",
        "message": "发生未知错误，请重试",
        "detail": "如果问题持续，请联系技术支持"
    }
}


def _get_error_response(error_type: str, extra_info: str = None) -> Dict[str, Any]:
    """
    获取用户友好的错误响应

    Args:
        error_type: 错误类型键值
        extra_info: 额外信息（可选）

    Returns:
        包含错误码、用户消息和技术详情的字典
    """
    error_info = ERROR_MESSAGES.get(error_type, ERROR_MESSAGES["unknown_error"])

    response = {
        "code": error_info["code"],
        "message": error_info["message"],
        "detail": error_info["detail"]
    }

    if extra_info:
        response["extra"] = extra_info

    return response


def _classify_error(error: Exception) -> str:
    """
    根据异常类型自动分类错误

    Args:
        error: 异常对象

    Returns:
        错误类型键值
    """
    error_str = str(error).lower()

    # DNS 解析失败 (Windows: 11001, Linux: -2)
    if "getaddrinfo failed" in error_str or "11001" in error_str or "name or service not known" in error_str:
        return "dns_resolution_failed"

    # 超时相关
    if "timeout" in error_str or "timed out" in error_str:
        return "llm_timeout"

    # 连接相关
    if "connection" in error_str or "connect" in error_str:
        if "refused" in error_str:
            return "llm_error"
        return "digital_human_connect_failed"

    # 语音相关
    if "speech" in error_str or "audio" in error_str or "asr" in error_str:
        return "speech_recognition_failed"

    if "tts" in error_str or "synthesize" in error_str:
        return "speech_synthesis_failed"

    # API 限流
    if "rate limit" in error_str or "429" in error_str:
        return "api_rate_limit"

    # 会话相关
    if "session" in error_str and ("not found" in error_str or "expired" in error_str):
        return "session_not_found"

    # 数据相关
    if "save" in error_str or "write" in error_str:
        return "data_save_failed"

    if "load" in error_str or "read" in error_str:
        return "data_load_failed"

    return "unknown_error"


async def _recognize_speech(audio_bytes: bytes, language: str = "zh") -> str:
    """使用 Sherpa ASR 进行语音识别"""
    try:
        from py.sherpa_asr import sherpa_recognize
        # 根据语言选择模型
        model_map = {
            "zh": "sherpa-onnx-sense-voice-zh-en-ja-ko-yue",
            "en": "sherpa-onnx-sense-voice-zh-en-ja-ko-yue",
            "ja": "sherpa-onnx-sense-voice-zh-en-ja-ko-yue",
            "ko": "sherpa-onnx-sense-voice-zh-en-ja-ko-yue",
            "yue": "sherpa-onnx-sense-voice-zh-en-ja-ko-yue"
        }
        model_name = model_map.get(language, "sherpa-onnx-sense-voice-zh-en-ja-ko-yue")
        text = await sherpa_recognize(audio_bytes, model_name)
        return text or ""
    except Exception as e:
        print(f"语音识别失败: {e}")
        raise RuntimeError(f"语音识别失败: {e}")


async def _synthesize_speech(
    text: str,
    language: str = "zh",
    engine: str = "edge",
    voice_id: str = "default"
) -> str:
    """使用 TTS 适配器合成语音，返回 base64 编码"""
    try:
        from py.tts_adapter import tts_adapter
        from py.get_setting import load_settings

        # 获取语音设置
        try:
            settings = await load_settings()
            engine = settings.get("tts_engine", engine)
            voice_id = settings.get("tts_voice", voice_id)
        except Exception:
            settings = {}

        # 合成音频
        audio_bytes = await tts_adapter.synthesize(
            text=text,
            engine=engine,
            voice_id=voice_id,
            language=language
        )

        # 返回 base64 编码
        return base64.b64encode(audio_bytes).decode('utf-8')

    except Exception as e:
        print(f"语音合成失败: {e}")
        # 返回空字符串而不是抛出异常，允许文本对话继续
        return ""


async def _get_llm_response(
    message: str,
    skill_id: str = None,
    session_id: str = None,
    history: List[Dict[str, str]] = None
) -> Dict[str, Any]:
    """调用 LLM 获取回复"""
    # 获取系统设置（从数据库加载）
    try:
        from py.get_setting import load_settings
        settings = await load_settings()
    except Exception:
        settings = {}

    # 构建消息
    messages = []

    # 添加技能系统提示词
    if skill_id and skill_id in SKILL_PROMPTS:
        messages.append({
            "role": "system",
            "content": SKILL_PROMPTS[skill_id]
        })
    else:
        messages.append({
            "role": "system",
            "content": """你是一位友好的研伴助手，帮助用户学习和研究。你的名字叫"研友"。
请用简洁、专业的语言回答问题，必要时提供学习建议。
由于是语音交互，请保持回复简洁，适合朗读。"""
        })

    # 添加历史消息（最多保留20条）
    if history:
        for msg in history[-20:]:
            messages.append({
                "role": msg.get("role", "user"),
                "content": msg.get("content", "")
            })

    # 添加用户消息
    messages.append({"role": "user", "content": message})

    # 获取 API 配置
    api_key = settings.get("api_key", "")
    base_url = settings.get("base_url", "https://api.openai.com/v1")
    model = settings.get("model", "gpt-3.5-turbo")

    if not api_key:
        # 模拟模式
        mock_response = "您好！我是研伴助手。\n\n⚠️ **未配置 API Key**\n\n要启用完整功能，请点击页面右上角的「⚙️ 设置」按钮配置语言模型 API。"
        return {
            "response": mock_response,
            "emotion": "neutral",
            "exp_gained": 5
        }

    # 调用 LLM API
    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                f"{base_url.rstrip('/')}/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": model,
                    "messages": messages,
                    "temperature": settings.get("temperature", 0.7),
                    "max_tokens": settings.get("max_tokens", 1024)  # 语音交互用较短的回复
                }
            )

            if response.status_code == 200:
                result = response.json()
                assistant_message = result.get("choices", [{}])[0].get("message", {}).get("content", "")

                # 分析情绪
                emotion = analyze_emotion(assistant_message)

                # 更新统计
                await update_chat_stats(skill_id)

                return {
                    "response": assistant_message,
                    "emotion": emotion,
                    "exp_gained": 10
                }
            else:
                error_detail = response.json().get("error", {}).get("message", "API 调用失败")
                raise Exception(error_detail)

    except httpx.TimeoutException:
        raise Exception("API 请求超时")
    except Exception as e:
        raise Exception(f"对话失败: {str(e)}")


async def _drive_tencent_digital_human(session_id: str, text: str) -> bool:
    """
    驱动腾讯数智人播报文本

    Args:
        session_id: 腾讯数智人会话 ID
        text: 要播报的文本

    Returns:
        是否成功驱动
    """
    try:
        from py.tencent_digital_human import client as tencent_client

        result = await tencent_client.drive_text(
            session_id=session_id,
            text=text,
            interrupt=True
        )

        # 检查是否成功
        if result.get("ErrCode") == 0:
            return True
        else:
            print(f"腾讯数智人驱动失败: {result.get('ErrMsg', '未知错误')}")
            return False

    except Exception as e:
        print(f"驱动腾讯数智人异常: {e}")
        return False


async def _save_voice_chat_history(
    session_id: str,
    user_message: str,
    assistant_message: str,
    skill_id: str = None
):
    """保存语音对话历史"""
    await ensure_storage()
    cache = await get_cache()

    # 保存到聊天历史
    data = await cache.get("chat_history", {"sessions": {}})

    data.setdefault("sessions", {})
    data["sessions"].setdefault(session_id, [])

    # 添加用户消息
    data["sessions"][session_id].append({
        "role": "user",
        "content": user_message,
        "timestamp": datetime.now().isoformat(),
        "type": "voice"
    })

    # 添加助手消息
    data["sessions"][session_id].append({
        "role": "assistant",
        "content": assistant_message,
        "timestamp": datetime.now().isoformat(),
        "type": "voice"
    })

    # 限制历史长度
    if len(data["sessions"][session_id]) > 100:
        data["sessions"][session_id] = data["sessions"][session_id][-100:]

    await cache.set("chat_history", data)

    # 同时记录协作贡献
    if skill_id:
        await _auto_record_collaboration(
            session_id=session_id,
            skill_id=skill_id,
            user_message=user_message,
            ai_message=assistant_message
        )


# ==================== 知识库管理 API ====================

@router.get("/knowledge/stats")
async def get_knowledge_stats():
    """获取知识库统计信息"""
    try:
        kb = await get_education_kb()
        stats = await kb.get_stats()
        return {"status": "success", "stats": stats}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.get("/knowledge/status")
async def get_knowledge_status():
    """
    获取知识库详细状态
    包括向量库、嵌入模型、BM25 等信息
    """
    try:
        kb = await get_education_kb()
        stats = await kb.get_stats()

        # 检测文件变更
        changes = await kb.detect_changes()

        return {
            "status": "success",
            "vector_store": {
                "type": "FAISS",
                "document_count": stats["document_count"],
                "index_size_mb": stats["index_size_mb"],
                "initialized": stats["initialized"],
                "last_updated": stats["last_init_time"]
            },
            "embedding": stats.get("embedding", {}),
            "config": stats.get("config", {}),
            "directories": {
                "kb_directory": stats.get("kb_directory"),
                "vector_directory": stats.get("vector_directory")
            },
            "pending_changes": {
                "added": len(changes["added"]),
                "removed": len(changes["removed"]),
                "modified": len(changes["modified"]),
                "has_changes": any(changes.values())
            }
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.post("/knowledge/preload")
async def preload_knowledge():
    """
    预加载知识库（服务启动后手动调用）
    避免首次对话时的延迟
    """
    try:
        success = await preload_education_kb()
        if success:
            kb = await get_education_kb()
            stats = await kb.get_stats()
            return {
                "status": "success",
                "message": "知识库预加载完成",
                "stats": stats
            }
        else:
            return {"status": "warning", "message": "知识库预加载失败，可能是嵌入模型未配置"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.post("/knowledge/search")
async def search_knowledge(query: str = Body(..., embed=True), k: int = Body(default=5, embed=True)):
    """
    搜索知识库

    Args:
        query: 搜索查询
        k: 返回结果数量

    Returns:
        检索结果列表
    """
    try:
        results = await search_education_knowledge(query, k=k)
        return {
            "status": "success",
            "query": query,
            "count": len(results),
            "results": results
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.post("/knowledge/rebuild")
async def rebuild_knowledge_base():
    """重建知识库向量索引"""
    try:
        kb = await get_education_kb()
        await kb.build_vector_store(force_rebuild=True)
        stats = await kb.get_stats()
        return {
            "status": "success",
            "message": "知识库向量索引重建完成",
            "stats": stats
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.post("/knowledge/incremental_update")
async def incremental_update_knowledge():
    """
    增量更新知识库
    只处理新增、删除、修改的文件，比完整重建更高效
    """
    try:
        kb = await get_education_kb()
        update_stats = await kb.incremental_update()
        stats = await kb.get_stats()
        return {
            "status": "success",
            "message": "知识库增量更新完成",
            "update_stats": update_stats,
            "current_stats": stats
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.get("/knowledge/changes")
async def detect_knowledge_changes():
    """
    检测知识库文件变更
    返回新增、删除、修改的文件列表
    """
    try:
        kb = await get_education_kb()
        changes = await kb.detect_changes()

        return {
            "status": "success",
            "changes": {
                "added": [{"path": str(p), "name": p.name} for p in changes["added"]],
                "removed": [{"path": str(p), "name": p.name} for p in changes["removed"]],
                "modified": [{"path": str(p), "name": p.name} for p in changes["modified"]]
            },
            "summary": {
                "added_count": len(changes["added"]),
                "removed_count": len(changes["removed"]),
                "modified_count": len(changes["modified"]),
                "total_changes": sum(len(v) for v in changes.values())
            }
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@router.get("/knowledge/files")
async def list_knowledge_files():
    """列出知识库中的所有文件"""
    from pathlib import Path

    files = []
    if EDUCATION_KB_DIR.exists():
        for file_path in EDUCATION_KB_DIR.glob("**/*"):
            if file_path.is_file() and file_path.suffix.lower() in [".md", ".txt"]:
                files.append({
                    "name": file_path.name,
                    "path": str(file_path.relative_to(EDUCATION_KB_DIR)),
                    "size": file_path.stat().st_size,
                    "modified": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
                })

    return {
        "status": "success",
        "directory": str(EDUCATION_KB_DIR),
        "count": len(files),
        "files": files
    }


# ==================== 导出功能 API ====================

from io import BytesIO
from fastapi.responses import StreamingResponse
from urllib.parse import quote

# 导出相关库
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[导出功能] python-docx 未安装，Word 导出不可用")


@router.get("/export/formats")
async def get_export_formats():
    """获取支持的导出格式"""
    formats = []
    if DOCX_AVAILABLE:
        formats.append({"id": "word", "name": "Word 文档", "extension": ".docx"})
    return {"formats": formats}


@router.post("/collaboration/export")
async def export_collaboration(request: ExportCollaborationRequest):
    """
    导出协作记录

    Args:
        request: 导出请求，包含 session_id, format, include_history

    Returns:
        文件下载流
    """
    try:
        session_id = request.session_id
        format = request.format
        include_history = request.include_history

        await ensure_storage()
        cache = await get_cache()

        # 获取协作数据
        data = await cache.get("collaboration", {
            "papers": [], "experiments": [], "reviews": [], "sessions": []
        })

        # 查找指定会话
        session = None
        session_type = None
        type_mapping = {
            "paper": "papers",
            "experiment": "experiments",
            "review": "reviews",
            "tutoring": "sessions"
        }

        for stype, skey in type_mapping.items():
            for s in data.get(skey, []):
                if s.get("id") == session_id:
                    session = s
                    session_type = stype
                    break
            if session:
                break

        if not session:
            raise HTTPException(status_code=404, detail="协作会话不存在")

        # 生成文档
        if not DOCX_AVAILABLE:
            raise HTTPException(status_code=400, detail="Word 导出功能不可用，请安装 python-docx")
        file_bytes = await _generate_collaboration_docx(session, session_type, include_history)
        filename = f"协作记录_{session_id}.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # 返回文件流
        encoded_filename = quote(filename)
        return StreamingResponse(
            BytesIO(file_bytes),
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"[导出错误] export_collaboration: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/collaboration/export_all")
async def export_all_collaborations(request: ExportAllCollaborationsRequest):
    """
    导出所有协作记录

    Args:
        request: 导出请求，包含 format, session_type, session_ids, include_chat_history

    Returns:
        文件下载流
    """
    try:
        session_type = request.session_type
        session_ids = request.session_ids
        include_chat_history = request.include_chat_history

        await ensure_storage()
        cache = await get_cache()

        data = await cache.get("collaboration", {
            "papers": [], "experiments": [], "reviews": [], "sessions": []
        })

        # 筛选会话
        all_sessions = []
        type_names = {
            "papers": "论文写作",
            "experiments": "实验设计",
            "reviews": "文献综述",
            "sessions": "虚拟辅导"
        }

        for skey, sname in type_names.items():
            if session_type:
                # 只导出指定类型
                stype_map = {"paper": "papers", "experiment": "experiments", "review": "reviews", "tutoring": "sessions"}
                if skey != stype_map.get(session_type):
                    continue

            for s in data.get(skey, []):
                # 如果指定了 session_ids，只导出这些会话
                if session_ids and s.get("id") not in session_ids:
                    continue
                s["_type_name"] = sname
                all_sessions.append(s)

        if not all_sessions:
            raise HTTPException(status_code=404, detail="没有可导出的协作记录")

        # 生成文档
        if not DOCX_AVAILABLE:
            raise HTTPException(status_code=400, detail="Word 导出功能不可用")
        file_bytes = await _generate_all_collaborations_docx(all_sessions, include_chat_history)
        filename = f"协作记录汇总_{datetime.now().strftime('%Y%m%d')}.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        encoded_filename = quote(filename)
        return StreamingResponse(
            BytesIO(file_bytes),
            media_type=media_type,
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"[导出错误] export_all_collaborations: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


async def _generate_collaboration_docx(session: dict, session_type: str, include_history: bool) -> bytes:
    """生成单个协作记录的 Word 文档"""
    doc = Document()

    # 标题
    type_names = {
        "paper": "论文写作协作",
        "experiment": "实验设计协作",
        "review": "文献综述协作",
        "tutoring": "虚拟导师辅导"
    }

    title = doc.add_heading(type_names.get(session_type, "协作记录"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 会话标题（如果有）
    if session.get("title"):
        p = doc.add_paragraph()
        p.add_run(session.get("title")).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 基本信息 - 简化表格
    doc.add_heading("基本信息", level=1)

    start_time = session.get("startTime", "")
    end_time = session.get("endTime", "进行中")

    # 计算时长
    duration_text = ""
    if start_time and end_time and end_time != "进行中":
        try:
            from datetime import datetime as dt
            start = dt.fromisoformat(start_time.replace("Z", "+00:00"))
            end = dt.fromisoformat(end_time.replace("Z", "+00:00"))
            duration_minutes = int((end - start).total_seconds() / 60)
            duration_text = f"{duration_minutes} 分钟"
        except:
            duration_text = session.get("duration", "")

    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ("开始时间", start_time[:19] if start_time else ""),
        ("结束时间", end_time[:19] if end_time and end_time != "进行中" else "进行中"),
        ("协作时长", duration_text or "未记录")
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # 协作统计 - 简化
    doc.add_heading("协作统计", level=1)

    ai_contribs = session.get("aiContributions", [])
    human_contribs = session.get("humanContributions", [])
    ai_count = len(ai_contribs)
    human_count = len(human_contribs)
    total = ai_count + human_count

    # 使用简洁的段落形式
    p = doc.add_paragraph()
    p.add_run(f"总交互次数: ").bold = True
    p.add_run(f"{total} 次  |  ")
    p.add_run(f"AI: ").bold = True
    p.add_run(f"{ai_count} 次  |  ")
    p.add_run(f"用户: ").bold = True
    p.add_run(f"{human_count} 次")

    if total > 0:
        p2 = doc.add_paragraph()
        p2.add_run(f"协作比例: AI {round(ai_count/total*100, 1)}% / 用户 {round(human_count/total*100, 1)}%")

    doc.add_paragraph()

    # 协作历史 - 优化格式
    if include_history and total > 0:
        doc.add_heading("对话记录", level=1)

        # 合并并按时间排序
        all_contributions = []
        for c in ai_contribs:
            all_contributions.append({**c, "role": "AI"})
        for c in human_contribs:
            all_contributions.append({**c, "role": "用户"})

        all_contributions.sort(key=lambda x: x.get("timestamp", ""))

        for i, c in enumerate(all_contributions, 1):
            role = c.get("role", "未知")
            content = c.get("content", "")
            timestamp = c.get("timestamp", "")

            # 格式化时间
            time_str = timestamp[11:16] if len(timestamp) > 16 else timestamp

            # 使用不同样式区分角色
            p = doc.add_paragraph()
            role_run = p.add_run(f"【{role}】")
            role_run.bold = True
            if role == "AI":
                role_run.font.color.rgb = None  # 默认颜色
            else:
                from docx.shared import RGBColor
                role_run.font.color.rgb = RGBColor(0, 100, 0)  # 绿色表示用户

            p.add_run(f" {time_str}")

            # 内容段落
            content_p = doc.add_paragraph(content)
            content_p.paragraph_format.left_indent = Pt(20)
            content_p.paragraph_format.space_after = Pt(8)

    # 摘要（如果有）
    if session.get("summary"):
        doc.add_heading("摘要", level=1)
        doc.add_paragraph(session.get("summary"))

    # 导出信息
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True

    # 保存到字节流
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


async def _generate_all_collaborations_docx(sessions: list, include_chat_history: bool = False) -> bytes:
    """生成所有协作记录汇总的 Word 文档"""
    try:
        doc = Document()

        # 标题
        title = doc.add_heading("协作记录汇总", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"共 {len(sessions)} 条协作记录")
        doc.add_paragraph()

        # 汇总统计 - 简化
        doc.add_heading("汇总统计", level=1)

        total_ai = sum(len(s.get("aiContributions", [])) for s in sessions)
        total_human = sum(len(s.get("humanContributions", [])) for s in sessions)
        total = total_ai + total_human

        # 使用简洁段落
        p = doc.add_paragraph()
        p.add_run(f"总会话数: ").bold = True
        p.add_run(f"{len(sessions)} 个  |  ")
        p.add_run(f"总交互: ").bold = True
        p.add_run(f"{total} 次  |  ")
        p.add_run(f"AI: ").bold = True
        p.add_run(f"{total_ai} 次  |  ")
        p.add_run(f"用户: ").bold = True
        p.add_run(f"{total_human} 次")

        doc.add_paragraph()

        # 各会话详情
        doc.add_heading("会话详情", level=1)

        for i, session in enumerate(sessions, 1):
            type_name = session.get("_type_name", "未知类型")
            title_text = session.get("title", "") or "无标题"
            start_time = session.get("startTime", "")
            end_time = session.get("endTime", "")

            # 会话标题
            p = doc.add_paragraph()
            p.add_run(f"{i}. [{type_name}] ").bold = True
            p.add_run(title_text)

            # 基本信息 - 简化
            time_str = start_time[:16] if start_time else ""
            end_str = end_time[:16] if end_time and end_time != "进行中" else "进行中"

            ai_count = len(session.get("aiContributions", []))
            human_count = len(session.get("humanContributions", []))

            info_p = doc.add_paragraph()
            info_p.add_run(f"   时间: {time_str} ~ {end_str}  |  交互: AI {ai_count}/用户 {human_count}")

            # 摘要（如果有）
            if session.get("summary"):
                summary_p = doc.add_paragraph()
                summary_p.add_run(f"   摘要: ").italic = True
                summary_p.add_run(session.get("summary"))

            # 聊天记录（如果请求）
            if include_chat_history and (ai_count + human_count) > 0:
                chat_p = doc.add_paragraph()
                chat_p.add_run("   对话记录:").italic = True

                # 合并并排序
                all_contribs = []
                for c in session.get("aiContributions", []):
                    all_contribs.append({**c, "role": "AI"})
                for c in session.get("humanContributions", []):
                    all_contribs.append({**c, "role": "用户"})
                all_contribs.sort(key=lambda x: x.get("timestamp", ""))

                for c in all_contribs:
                    role = c.get("role", "")
                    content = c.get("content", "")
                    ts = c.get("timestamp", "")
                    time_short = ts[11:16] if len(ts) > 16 else ts

                    # 缩进的对话行
                    chat_line = doc.add_paragraph()
                    chat_line.paragraph_format.left_indent = Pt(30)
                    chat_line.add_run(f"[{role}] {time_short}: ").bold = (role == "用户")
                    # 截断过长的内容
                    if len(content) > 500:
                        chat_line.add_run(content[:500] + "...")
                    else:
                        chat_line.add_run(content)

            doc.add_paragraph()  # 会话间空行

        # 页脚
        p = doc.add_paragraph()
        p.add_run("由研伴系统生成").italic = True

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"[导出错误] _generate_all_collaborations_docx: {e}")
        import traceback
        traceback.print_exc()
        raise


# ==================== 学习报告生成 API ====================

@router.get("/report/generate")
async def generate_learning_report():
    """
    生成学习报告

    汇总成长数据、协作统计、技能使用、成就等信息
    """
    try:
        await ensure_storage()
        cache = await get_cache()

        # 获取各类数据
        growth_data = await cache.get("growth", {
            "level": 1, "exp": 0, "totalExp": 0, "achievements": [], "stats": {}, "skillUsage": {}
        })

        collab_data = await cache.get("collaboration", {
            "papers": [], "experiments": [], "reviews": [], "sessions": []
        })

        achievement_data = await cache.get("achievement", {"unlocked": [], "history": []})

        # 计算统计数据
        total_sessions = sum(len(collab_data.get(k, [])) for k in ["papers", "experiments", "reviews", "sessions"])

        total_ai_contrib = 0
        total_human_contrib = 0
        for key in ["papers", "experiments", "reviews", "sessions"]:
            for session in collab_data.get(key, []):
                total_ai_contrib += len(session.get("aiContributions", []))
                total_human_contrib += len(session.get("humanContributions", []))

        # 技能使用统计
        skill_usage = growth_data.get("skillUsage", {})
        skill_names = {
            "research-assistant": "科研助手",
            "literature-review": "文献综述",
            "paper-writing": "论文写作",
            "academic-tutoring": "虚拟导师"
        }

        # 成就统计
        unlocked_count = len(achievement_data.get("unlocked", []))

        # 构建报告
        report = {
            "generatedAt": datetime.now().isoformat(),
            "summary": {
                "level": growth_data.get("level", 1),
                "totalExp": growth_data.get("totalExp", 0),
                "currentExp": growth_data.get("exp", 0),
                "expForNextLevel": growth_data.get("level", 1) * 100
            },
            "stats": growth_data.get("stats", {}),
            "collaboration": {
                "totalSessions": total_sessions,
                "aiContributions": total_ai_contrib,
                "humanContributions": total_human_contrib,
                "collaborationRatio": {
                    "ai": round(total_ai_contrib / (total_ai_contrib + total_human_contrib) * 100, 1) if (total_ai_contrib + total_human_contrib) > 0 else 0,
                    "human": round(total_human_contrib / (total_ai_contrib + total_human_contrib) * 100, 1) if (total_ai_contrib + total_human_contrib) > 0 else 0
                }
            },
            "skills": {
                "usage": {skill_names.get(k, k): v for k, v in skill_usage.items()},
                "mostUsed": max(skill_usage.items(), key=lambda x: x[1])[0] if skill_usage else None
            },
            "achievements": {
                "unlockedCount": unlocked_count,
                "recentUnlocks": achievement_data.get("history", [])[-5:]  # 最近5个成就
            },
            "recommendations": _generate_recommendations(growth_data, skill_usage, total_sessions)
        }

        return report

    except Exception as e:
        print(f"[教育API] 生成学习报告失败: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"生成学习报告失败: {str(e)}")


def _generate_recommendations(growth_data: dict, skill_usage: dict, total_sessions: int) -> list:
    """生成学习建议"""
    recommendations = []

    stats = growth_data.get("stats", {})

    # 基于对话次数
    conversations = stats.get("conversations", 0)
    if conversations < 10:
        recommendations.append("建议多与数字人互动，探索不同的教育技能")
    elif conversations > 100:
        recommendations.append("您是活跃用户！可以尝试更高级的科研协作功能")

    # 基于技能使用
    if skill_usage:
        used_skills = set(skill_usage.keys())
        all_skills = {"research-assistant", "literature-review", "paper-writing", "academic-tutoring"}
        unused_skills = all_skills - used_skills

        if unused_skills:
            skill_names = {
                "research-assistant": "科研助手",
                "literature-review": "文献综述",
                "paper-writing": "论文写作",
                "academic-tutoring": "虚拟导师"
            }
            unused_names = [skill_names.get(s, s) for s in unused_skills]
            recommendations.append(f"建议尝试使用: {', '.join(unused_names)}")

    # 基于协作次数
    if total_sessions < 5:
        recommendations.append("开始您的第一个协作项目，记录学习过程")
    elif total_sessions > 20:
        recommendations.append("协作经验丰富！可以考虑导出学习报告进行总结")

    # 基于等级
    level = growth_data.get("level", 1)
    if level < 3:
        recommendations.append("继续积累经验，解锁更多成就")
    elif level >= 10:
        recommendations.append("您已是资深学者！欢迎分享您的学习经验")

    return recommendations if recommendations else ["继续保持学习热情！"]


@router.post("/report/export")
async def export_learning_report(request: ExportReportRequest):
    """
    导出学习报告

    Args:
        request: 导出请求，包含 format

    Returns:
        文件下载流
    """
    # 获取报告数据
    report = await generate_learning_report()

    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=400, detail="Word 导出功能不可用")
    file_bytes = await _generate_report_docx(report)
    filename = f"学习报告_{datetime.now().strftime('%Y%m%d')}.docx"
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    encoded_filename = quote(filename)
    return StreamingResponse(
        BytesIO(file_bytes),
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )


async def _generate_report_docx(report: dict) -> bytes:
    """生成学习报告 Word 文档"""
    doc = Document()

    # 标题
    title = doc.add_heading("学习报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"生成时间: {report.get('generatedAt', '')}")
    doc.add_paragraph()

    # 成长概览
    doc.add_heading("成长概览", level=1)

    summary = report.get("summary", {})
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'

    summary_data = [
        ("当前等级", f"Lv.{summary.get('level', 1)}"),
        ("累计经验", f"{summary.get('totalExp', 0)} 点"),
        ("当前经验", f"{summary.get('currentExp', 0)} / {summary.get('expForNextLevel', 100)} 点"),
        ("升级进度", f"{round(summary.get('currentExp', 0) / summary.get('expForNextLevel', 100) * 100, 1)}%")
    ]

    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 学习统计
    doc.add_heading("学习统计", level=1)

    stats = report.get("stats", {})
    stats_table = doc.add_table(rows=6, cols=2)
    stats_table.style = 'Table Grid'

    stats_data = [
        ("对话次数", str(stats.get("conversations", 0))),
        ("阅读文献", f"{stats.get('papers_read', 0)} 篇"),
        ("设计实验", f"{stats.get('experiments_designed', 0)} 个"),
        ("论文写作", f"{stats.get('papers_written', 0)} 篇"),
        ("辅导会话", f"{stats.get('tutoring_sessions', 0)} 次"),
        ("学习时长", f"{round(stats.get('hours_spent', 0), 1)} 小时")
    ]

    for i, (label, value) in enumerate(stats_data):
        stats_table.rows[i].cells[0].text = label
        stats_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 协作统计
    doc.add_heading("协作统计", level=1)

    collab = report.get("collaboration", {})
    collab_table = doc.add_table(rows=4, cols=2)
    collab_table.style = 'Table Grid'

    ratio = collab.get("collaborationRatio", {})
    collab_data = [
        ("协作会话", f"{collab.get('totalSessions', 0)} 次"),
        ("AI 贡献", f"{collab.get('aiContributions', 0)} 次 ({ratio.get('ai', 0)}%)"),
        ("用户贡献", f"{collab.get('humanContributions', 0)} 次 ({ratio.get('human', 0)}%)"),
        ("协作比例", f"AI {ratio.get('ai', 0)}% : 用户 {ratio.get('human', 0)}%")
    ]

    for i, (label, value) in enumerate(collab_data):
        collab_table.rows[i].cells[0].text = label
        collab_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 技能使用
    doc.add_heading("技能使用", level=1)

    skills = report.get("skills", {})
    usage = skills.get("usage", {})

    if usage:
        skills_table = doc.add_table(rows=len(usage) + 1, cols=2)
        skills_table.style = 'Table Grid'
        skills_table.rows[0].cells[0].text = "技能名称"
        skills_table.rows[0].cells[1].text = "使用次数"

        for i, (skill_name, count) in enumerate(usage.items(), 1):
            skills_table.rows[i].cells[0].text = skill_name
            skills_table.rows[i].cells[1].text = str(count)
    else:
        doc.add_paragraph("暂无技能使用记录")

    doc.add_paragraph()

    # 成就
    doc.add_heading("成就", level=1)

    achievements = report.get("achievements", {})
    doc.add_paragraph(f"已解锁成就: {achievements.get('unlockedCount', 0)} 个")

    recent = achievements.get("recentUnlocks", [])
    if recent:
        doc.add_paragraph("最近解锁:")
        for ach in recent:
            doc.add_paragraph(f"  - {ach.get('id', '')} ({ach.get('unlockedAt', '')})")

    doc.add_paragraph()

    # 学习建议
    doc.add_heading("学习建议", level=1)

    recommendations = report.get("recommendations", [])
    for rec in recommendations:
        doc.add_paragraph(f"• {rec}")

    # 页脚
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("由研伴系统生成").italic = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==================== 智能笔记系统 ====================

# 笔记数据文件
NOTES_DATA_FILE = EDUCATION_DATA_DIR / "notes_data.json"

class NoteModel(BaseModel):
    """笔记数据模型"""
    id: str
    title: str
    content: str
    summary: Optional[str] = None
    tags: List[str] = []
    skill_id: Optional[str] = None
    skill_name: Optional[str] = None
    source_type: str = "manual"  # manual, ai_generated, from_chat
    source_session_id: Optional[str] = None
    related_notes: List[str] = []
    created_at: str
    updated_at: str

class NoteCreateRequest(BaseModel):
    """创建笔记请求"""
    title: str
    content: str
    tags: List[str] = []
    skill_id: Optional[str] = None
    skill_name: Optional[str] = None
    source_type: str = "manual"
    source_session_id: Optional[str] = None

class NoteUpdateRequest(BaseModel):
    """更新笔记请求"""
    title: Optional[str] = None
    content: Optional[str] = None
    tags: Optional[List[str]] = None

class AIGenerateNoteRequest(BaseModel):
    """AI生成笔记请求"""
    session_id: Optional[str] = None
    messages: Optional[List[Dict[str, str]]] = None
    skill_id: Optional[str] = None
    skill_name: Optional[str] = None
    custom_prompt: Optional[str] = None


async def load_notes() -> Dict[str, Any]:
    """加载笔记数据"""
    if NOTES_DATA_FILE.exists():
        try:
            with open(NOTES_DATA_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return {"notes": [], "tags": {}}


async def save_notes(data: Dict[str, Any]):
    """保存笔记数据"""
    with open(NOTES_DATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


@router.get("/notes")
async def get_notes(
    skill_id: Optional[str] = None,
    tag: Optional[str] = None,
    search: Optional[str] = None,
    limit: int = 20,
    offset: int = 0
):
    """获取笔记列表"""
    data = await load_notes()
    notes = data.get("notes", [])

    # 过滤
    if skill_id:
        notes = [n for n in notes if n.get("skill_id") == skill_id]
    if tag:
        notes = [n for n in notes if tag in n.get("tags", [])]
    if search:
        search_lower = search.lower()
        notes = [n for n in notes if
                 search_lower in n.get("title", "").lower() or
                 search_lower in n.get("content", "").lower() or
                 search_lower in n.get("summary", "").lower()]

    # 按更新时间排序
    notes.sort(key=lambda x: x.get("updated_at", ""), reverse=True)

    total = len(notes)
    notes = notes[offset:offset + limit]

    return {
        "notes": notes,
        "total": total,
        "tags": data.get("tags", {})
    }


@router.post("/notes")
async def create_note(request: NoteCreateRequest = Body(...)):
    """创建笔记"""
    import uuid
    data = await load_notes()

    now = datetime.now().isoformat()
    note = {
        "id": str(uuid.uuid4()),
        "title": request.title,
        "content": request.content,
        "summary": request.content[:200] + "..." if len(request.content) > 200 else request.content,
        "tags": request.tags,
        "skill_id": request.skill_id,
        "skill_name": request.skill_name,
        "source_type": request.source_type,
        "source_session_id": request.source_session_id,
        "related_notes": [],
        "created_at": now,
        "updated_at": now
    }

    data["notes"].append(note)

    # 更新标签统计
    tags = data.get("tags", {})
    for tag in request.tags:
        tags[tag] = tags.get(tag, 0) + 1
    data["tags"] = tags

    await save_notes(data)

    # 更新成长统计
    cache = await get_cache()
    await cache.increment("growth", "stats.notesCreated", 1)
    await cache.increment("growth", "exp", 10)
    await cache.increment("growth", "totalExp", 10)

    return {"status": "success", "note": note}


@router.get("/notes/{note_id}")
async def get_note(note_id: str):
    """获取单个笔记详情"""
    data = await load_notes()
    notes = data.get("notes", [])

    for note in notes:
        if note.get("id") == note_id:
            return note

    raise HTTPException(status_code=404, detail="笔记不存在")


@router.put("/notes/{note_id}")
async def update_note(note_id: str, request: NoteUpdateRequest = Body(...)):
    """更新笔记"""
    data = await load_notes()
    notes = data.get("notes", [])

    for i, note in enumerate(notes):
        if note.get("id") == note_id:
            if request.title is not None:
                note["title"] = request.title
            if request.content is not None:
                note["content"] = request.content
                note["summary"] = request.content[:200] + "..." if len(request.content) > 200 else request.content
            if request.tags is not None:
                # 更新标签统计
                old_tags = set(note.get("tags", []))
                new_tags = set(request.tags)
                tags = data.get("tags", {})

                for tag in old_tags - new_tags:
                    if tag in tags and tags[tag] > 0:
                        tags[tag] -= 1
                for tag in new_tags - old_tags:
                    tags[tag] = tags.get(tag, 0) + 1

                note["tags"] = request.tags
                data["tags"] = tags

            note["updated_at"] = datetime.now().isoformat()
            notes[i] = note
            data["notes"] = notes
            await save_notes(data)
            return {"status": "success", "note": note}

    raise HTTPException(status_code=404, detail="笔记不存在")


@router.delete("/notes/{note_id}")
async def delete_note(note_id: str):
    """删除笔记"""
    data = await load_notes()
    notes = data.get("notes", [])

    for i, note in enumerate(notes):
        if note.get("id") == note_id:
            # 更新标签统计
            tags = data.get("tags", {})
            for tag in note.get("tags", []):
                if tag in tags and tags[tag] > 0:
                    tags[tag] -= 1

            notes.pop(i)
            data["notes"] = notes
            data["tags"] = tags
            await save_notes(data)
            return {"status": "success"}

    raise HTTPException(status_code=404, detail="笔记不存在")


@router.post("/notes/ai_generate")
async def ai_generate_note(request: AIGenerateNoteRequest = Body(...)):
    """AI自动生成笔记"""
    import uuid
    import re
    import time as time_module
    from collections import Counter

    start_time = time_module.time()
    print("[笔记生成] 开始处理请求")

    # 获取对话内容
    messages = request.messages or []

    if not messages and request.session_id:
        # 从历史记录获取消息
        if CHAT_HISTORY_FILE.exists():
            try:
                with open(CHAT_HISTORY_FILE, 'r', encoding='utf-8') as f:
                    history = json.load(f)
                for session in history.get("sessions", []):
                    if session.get("session_id") == request.session_id:
                        messages = session.get("messages", [])[-20:]  # 最近20条
                        break
            except Exception:
                pass

    if not messages:
        raise HTTPException(status_code=400, detail="没有可用的对话内容")

    print(f"[笔记生成] 获取到 {len(messages)} 条消息, 耗时: {time_module.time() - start_time:.2f}s")

    # 构建AI提示词
    conversation_text = "\n".join([
        f"{'用户' if m.get('role') == 'user' else 'AI'}: {m.get('content', '')}"
        for m in messages
    ])

    system_prompt = """你是一个学习笔记助手。请根据以下对话内容，生成一份结构化的学习笔记。

要求：
1. 提取核心知识点和概念
2. 使用Markdown格式，包含标题、列表、重点标记
3. 总结关键要点
4. 如果有方法论或步骤，请清晰列出
5. 标注可能的疑问点或需要深入学习的方向

输出格式：
## [主题标题]

### 核心概念
- ...

### 关键要点
1. ...

### 方法/步骤
1. ...

### 疑问与深入方向
- ..."""

    # 调用AI生成
    t_ai_start = time_module.time()
    client = get_global_http_client()
    settings = await get_cached_settings()

    ai_messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"请根据以下对话生成学习笔记：\n\n{conversation_text}"}
    ]

    try:
        base_url = settings.get("base_url", "https://api.openai.com/v1")
        api_key = settings.get("api_key", "")
        model = settings.get("model", "gpt-4o-mini")

        # 优先使用快速模型配置（笔记生成不需要复杂模型）
        fast_settings = settings.get("fast", {})
        model_providers = settings.get("modelProviders", [])
        fast_provider_id = fast_settings.get("selectedProvider")

        if fast_provider_id and model_providers:
            for provider in model_providers:
                if str(provider.get("id", "")) == str(fast_provider_id):
                    if provider.get("apiKey"):
                        base_url = provider.get("baseUrl", base_url)
                        api_key = provider.get("apiKey", api_key)
                        model = provider.get("model", model)
                        print(f"[笔记生成] 使用快速模型配置: {model}")
                        break

        print(f"[笔记生成] 调用AI: model={model}, base_url={base_url}")

        response = await client.post(
            f"{base_url}/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            json={
                "model": model,
                "messages": ai_messages,
                "max_tokens": 2000,
                "temperature": 0.7
            }
        )

        print(f"[笔记生成] AI响应完成, 耗时: {time_module.time() - t_ai_start:.2f}s, 状态码: {response.status_code}")

        if response.status_code != 200:
            error_text = response.text[:200] if response.text else "无错误信息"
            print(f"[笔记生成] AI调用失败: {error_text}")
            raise HTTPException(status_code=500, detail=f"AI生成失败: {error_text}")

        result = response.json()
        content = result["choices"][0]["message"]["content"]

        # 提取标题（第一个##后的内容）
        title_match = re.search(r'##\s*(.+)', content)
        title = title_match.group(1) if title_match else "学习笔记"

        # 生成摘要
        summary = content[:200] + "..." if len(content) > 200 else content

        # 自动生成标签
        tags = []
        if request.skill_name:
            tags.append(request.skill_name)
        # 简单的关键词提取
        keywords = re.findall(r'[\u4e00-\u9fa5]{2,4}', content)
        tag_candidates = [k for k, _ in Counter(keywords).most_common(3)]
        tags.extend(tag_candidates)

        # 保存笔记
        data = await load_notes()
        now = datetime.now().isoformat()
        note = {
            "id": str(uuid.uuid4()),
            "title": title,
            "content": content,
            "summary": summary,
            "tags": list(set(tags)),
            "skill_id": request.skill_id,
            "skill_name": request.skill_name,
            "source_type": "ai_generated",
            "source_session_id": request.session_id,
            "related_notes": [],
            "created_at": now,
            "updated_at": now
        }

        data["notes"].append(note)

        # 更新标签统计
        note_tags = data.get("tags", {})
        for tag in note["tags"]:
            note_tags[tag] = note_tags.get(tag, 0) + 1
        data["tags"] = note_tags

        await save_notes(data)

        # 更新成长统计
        cache = await get_cache()
        await cache.increment("growth", "stats.notesCreated", 1)
        await cache.increment("growth", "exp", 15)
        await cache.increment("growth", "totalExp", 15)

        return {"status": "success", "note": note}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI生成失败: {str(e)}")


@router.get("/notes/{note_id}/related")
async def get_related_notes(note_id: str, limit: int = 5):
    """获取相关笔记推荐"""
    data = await load_notes()
    notes = data.get("notes", [])

    # 找到当前笔记
    current_note = None
    for note in notes:
        if note.get("id") == note_id:
            current_note = note
            break

    if not current_note:
        raise HTTPException(status_code=404, detail="笔记不存在")

    # 简单的相关性计算：基于标签和技能
    related = []
    current_tags = set(current_note.get("tags", []))
    current_skill = current_note.get("skill_id")

    for note in notes:
        if note.get("id") == note_id:
            continue

        score = 0
        note_tags = set(note.get("tags", []))

        # 标签重叠度
        tag_overlap = len(current_tags & note_tags)
        score += tag_overlap * 2

        # 同技能
        if current_skill and note.get("skill_id") == current_skill:
            score += 3

        if score > 0:
            related.append({"note": note, "score": score})

    # 按相关度排序
    related.sort(key=lambda x: x["score"], reverse=True)
    related = [r["note"] for r in related[:limit]]

    return {"related_notes": related}


@router.post("/notes/export")
async def export_notes(note_ids: List[str] = Body(...)):
    """导出笔记为Word文档"""
    from docx import Document
    from io import BytesIO
    import re

    def add_formatted_runs(paragraph, text):
        """将Markdown行内格式转为Word格式（加粗、斜体、代码）"""
        # 匹配 **加粗**、*斜体*、`代码` 的混合文本
        pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
        last_end = 0

        for match in re.finditer(pattern, text):
            # 匹配前的普通文本
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            full = match.group(0)
            if full.startswith("**"):
                run = paragraph.add_run(match.group(2))
                run.bold = True
            elif full.startswith("*"):
                run = paragraph.add_run(match.group(3))
                run.italic = True
            elif full.startswith("`"):
                run = paragraph.add_run(match.group(4))
                run.font.name = "宋体"
                run.font.size = Pt(10)

            last_end = match.end()

        # 剩余普通文本
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    data = await load_notes()
    notes = data.get("notes", [])

    # 筛选要导出的笔记
    export_list = [n for n in notes if n.get("id") in note_ids]

    if not export_list:
        raise HTTPException(status_code=400, detail="没有可导出的笔记")

    doc = Document()
    doc.add_heading("研伴学习笔记", 0)

    for note in export_list:
        doc.add_heading(note.get("title", "无标题"), level=1)

        # 元信息
        meta = doc.add_paragraph()
        meta.add_run(f"创建时间: {note.get('created_at', '')}").italic = True
        if note.get("skill_name"):
            meta.add_run(f" | 技能: {note.get('skill_name')}").italic = True
        if note.get("tags"):
            meta.add_run(f" | 标签: {', '.join(note.get('tags', []))}").italic = True

        doc.add_paragraph()

        # 内容 - 完整的Markdown转Word
        content = note.get("content", "")
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            # 标题
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            # 无序列表
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_runs(p, stripped[2:])
            # 有序列表
            elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == '.':
                p = doc.add_paragraph(style='List Number')
                add_formatted_runs(p, stripped[3:])
            # 普通段落
            else:
                p = doc.add_paragraph()
                add_formatted_runs(p, stripped)

        doc.add_paragraph()
        doc.add_paragraph("─" * 40)
        doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=notes_export.docx"}
    )
